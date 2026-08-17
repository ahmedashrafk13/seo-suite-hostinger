#!/usr/bin/env python3
"""
Word (.docx) report writer for the Internal Linking Agent.
=========================================================

This is the ONLY report generator. There is deliberately no parallel Markdown
builder: two renderers over the same numbers is how a deliverable ends up
disagreeing with itself, and that failure mode has already been paid for once on
this project (an "editorial link" count derived two different ways reported 316
links / 54 orphans when the truth was 52 / 171).

Everything written here comes from values computed by internal_link_agent.py and
is passed in; this module derives no metrics of its own beyond formatting
arithmetic (percentages, per-100-word rates) that is shown alongside its inputs.

Design goals for the document itself
------------------------------------
  * It has to survive being read by someone who did not run it. Cover page,
    executive summary with the three numbers that matter, then evidence.
  * Every claim is traceable. Recommendations carry the exact sentence and the
    character offset; every finding names the pages it came from.
  * Limitations are in the body, not a footnote. A report that hides its coverage
    gaps is worse than no report.
  * Colour is used for severity only, never for decoration, and every colour is
    paired with a word so the document still works printed in greyscale.
"""

from __future__ import annotations

from datetime import datetime, timezone

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Palette - dark ink on white, one accent, three severity colours.
# --------------------------------------------------------------------------- #

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5F, 0x6B, 0x7A)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
OK = RGBColor(0x1E, 0x7B, 0x34)
WARN = RGBColor(0xA8, 0x5C, 0x00)
BAD = RGBColor(0xB3, 0x1B, 0x1B)

SHADE_HEADER = "1F4E79"
SHADE_ROW = "F2F5F8"
SHADE_CALLOUT = "FFF6E5"
SHADE_NOTE = "EEF3F8"

SEVERITY_COLOR = {"critical": BAD, "high": BAD, "medium": WARN, "low": MUTED}
TIER_COLOR = {"high": OK, "single-word": WARN, "needs-new-sentence": MUTED}

MAX_URL = 74          # URLs are truncated for table fit; CSVs hold the full value
TABLE_STYLE = "Table Grid"


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #

def _shade(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_text(cell, text: str, *, bold=False, size=8.5, color=None,
               align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color if color is not None else INK


def _no_autofit(table, widths: list[float]) -> None:
    """
    Pin column widths. Word ignores a table's own layout hint unless autofit is
    off AND every cell in the column carries the width, which is why this loops
    over rows rather than setting table.columns[i].width once.
    """
    table.autofit = False
    try:
        table.allow_autofit = False
    except Exception:
        pass
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def _keep_header_repeating(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _truncate(text: str, limit: int) -> str:
    text = str(text or "")
    return text if len(text) <= limit else text[: limit - 1] + "…"


# --------------------------------------------------------------------------- #
# Report builder
# --------------------------------------------------------------------------- #

class Report:
    def __init__(self, site: str):
        self.doc = Document()
        self.site = site
        self._setup_styles()
        self._section_no = 0

    # ---- setup --------------------------------------------------------- #

    def _setup_styles(self) -> None:
        doc = self.doc
        for s in doc.sections:
            s.top_margin = Inches(0.85)
            s.bottom_margin = Inches(0.85)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
        # Latin font alone leaves East-Asian runs on the theme font; set both so
        # the document renders consistently everywhere.
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rfonts.set(qn("w:eastAsia"), "Calibri")

        for name, size, color, before, after in (
            ("Heading 1", 17, ACCENT, 16, 8),
            ("Heading 2", 13, ACCENT, 13, 6),
            ("Heading 3", 11.5, INK, 10, 4),
            ("Heading 4", 10.5, MUTED, 8, 3),
        ):
            st = doc.styles[name]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.color.rgb = color
            st.font.italic = False
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
            st.paragraph_format.keep_with_next = True

    # ---- primitives ---------------------------------------------------- #

    def para(self, text: str = "", *, size=10.5, bold=False, italic=False,
             color=None, align=None, space_after=6, indent=0.0):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = color if color is not None else INK
        return p

    def rich(self, parts: list[tuple[str, dict]], *, size=10.5, space_after=6,
             indent=0.0):
        """A paragraph from (text, formatting) pairs, for inline emphasis."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        for text, fmt in parts:
            run = p.add_run(text)
            run.bold = fmt.get("bold", False)
            run.italic = fmt.get("italic", False)
            run.font.size = Pt(fmt.get("size", size))
            run.font.color.rgb = fmt.get("color", INK)
            if fmt.get("mono"):
                run.font.name = "Consolas"
            if fmt.get("highlight"):
                from docx.enum.text import WD_COLOR_INDEX
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return p

    def bullet(self, text: str, *, bold_prefix: str = "", level=0):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.28 + 0.22 * level)
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            r.font.size = Pt(10.5)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        return p

    def h1(self, text: str, numbered=True):
        self.doc.add_page_break()
        if numbered:
            self._section_no += 1
            text = f"{self._section_no}. {text}"
        return self.doc.add_heading(text, level=1)

    def h2(self, text: str):
        return self.doc.add_heading(text, level=2)

    def h3(self, text: str):
        return self.doc.add_heading(text, level=3)

    def callout(self, title: str, body: str, *, kind="note"):
        """A single-cell shaded table. Word has no native callout box."""
        fill = {"note": SHADE_NOTE, "warn": SHADE_CALLOUT}.get(kind, SHADE_NOTE)
        color = {"note": ACCENT, "warn": WARN}.get(kind, ACCENT)
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        _shade(cell, fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = color
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(body)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = INK
        self.para(space_after=4)
        return t

    def table(self, headers: list[str], rows: list[list], widths: list[float],
              *, aligns: list | None = None, row_colors: list | None = None,
              size=8.5):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = TABLE_STYLE
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0]
        for i, h in enumerate(headers):
            _shade(hdr.cells[i], SHADE_HEADER)
            _cell_text(hdr.cells[i], h, bold=True, size=size,
                       color=RGBColor(0xFF, 0xFF, 0xFF))
        _keep_header_repeating(hdr)
        for r_i, row in enumerate(rows):
            cells = t.add_row().cells
            for c_i, val in enumerate(row):
                if r_i % 2 == 1:
                    _shade(cells[c_i], SHADE_ROW)
                col = None
                if row_colors and r_i < len(row_colors):
                    rc = row_colors[r_i]
                    if isinstance(rc, dict):
                        col = rc.get(c_i)
                    elif c_i == 0:
                        col = rc
                _cell_text(cells[c_i], val, size=size, color=col,
                           align=(aligns[c_i] if aligns and c_i < len(aligns) else None))
        _no_autofit(t, widths)
        self.para(space_after=4)
        return t

    def kv_table(self, rows: list[tuple[str, str, object]]):
        """Metric / value pairs. Third element optionally colours the value."""
        t = self.doc.add_table(rows=0, cols=2)
        t.style = TABLE_STYLE
        for label, value, color in rows:
            cells = t.add_row().cells
            _cell_text(cells[0], label, size=9.5)
            _cell_text(cells[1], value, bold=True, size=9.5,
                       color=color if color else INK,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
        _no_autofit(t, [4.9, 1.9])
        self.para(space_after=4)
        return t

    def toc_field(self) -> None:
        """
        A real Word TOC field. It renders as a placeholder line until the reader
        presses F9 / opens the document in Word, so the placeholder text says so
        rather than looking like a rendering bug.
        """
        p = self.doc.add_paragraph()
        run = p.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-2" \h \z \u'
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = ("Contents will appear here - right-click and choose "
                            "“Update Field” in Word.")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (fld, instr, sep, placeholder, end):
            run._r.append(el)

    def page_footer(self) -> None:
        """Site + page number in the footer of every page."""
        footer = self.doc.sections[0].footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Internal Linking Audit — {self.site}   |   page ")
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED
        run = p.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    def save(self, path) -> None:
        self.doc.save(str(path))


# --------------------------------------------------------------------------- #
# Document content
# --------------------------------------------------------------------------- #

def _sev_color(sev: str):
    return SEVERITY_COLOR.get(str(sev).lower(), MUTED)


def build_docx(path, *, root, pages, urls, graph, recs, orphans, underlinked,
               cannibal, saturated, crawler, cfg, elapsed, coverage, broken,
               malformed, dup_clusters, empty_pages, non_content, link_heavy,
               reject_stats, gsc=None, summary=None) -> None:
    """
    Render the full audit as a Word document.

    Every argument is a value already computed by the pipeline. Nothing is
    recomputed here, so the document cannot disagree with summary.json.
    """
    n = len(urls)
    high = [r for r in recs if r["confidence"] == "high"]
    single = [r for r in recs if r["confidence"] == "single-word"]
    needs = [r for r in recs if r["confidence"] == "needs-new-sentence"]
    dup_page_set = {u for cl in dup_clusters for u in cl}
    total_ed = len(graph["editorial_edges"])
    # `needs-verification` is a duplicate-shaped finding whose EXTRACTION could
    # not be trusted (both pages fell back to structural-only parsing, which is
    # what a JavaScript-rendered site looks like to the crawler). It belongs in
    # the duplicate-content discussion with a caveat — not in "competing pages",
    # which describes a genuine topic overlap and prescribes a different fix.
    dupes = [c for c in cannibal if c["severity"] == "critical"]
    unverified_dupes = [c for c in cannibal if c["severity"] == "needs-verification"]
    rivals = [c for c in cannibal
              if c["severity"] not in ("critical", "needs-verification")]
    real_broken = [b for b in broken
                   if b["referring_pages"] and b.get("classification") == "broken link"]
    transient = [b for b in broken if b.get("classification") != "broken link"]
    now = datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC")

    R = Report(root)
    R.page_footer()

    # ===================== COVER ========================================= #
    R.para(space_after=52)
    R.para("INTERNAL LINKING AUDIT", size=27, bold=True, color=ACCENT,
           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    R.para(root, size=15, color=INK, space_after=26)
    R.para("Crawl-based analysis of internal link structure, orphan pages, anchor "
           "text opportunities and keyword competition.",
           size=11.5, color=MUTED, space_after=30)

    R.kv_table([
        ("Site audited", root, None),
        ("Report generated", now, None),
        ("Pages analysed", f"{n}", ACCENT),
        ("Crawl duration", f"{elapsed:.0f} seconds", None),
        ("Link recommendations", f"{len(recs)}", ACCENT),
        ("Ready to implement today", f"{len(high)}", OK if high else MUTED),
        ("Orphan pages found", f"{len(orphans)}", BAD if orphans else OK),
        ("Analysis method", "Deterministic TF-IDF - no language model", None),
    ])

    R.para(space_after=18)
    R.callout(
        "How to read this report",
        "Section 2 is the executive summary. Section 4 is the work queue: each "
        "entry names a source page, a target page, and an anchor phrase that "
        "already exists word-for-word in the source page's own copy, quoted with "
        "the sentence around it so it can be located and checked in seconds. "
        "Nothing in this document is paraphrased or invented by a language model. "
        "Every table is also delivered as a CSV alongside this file.")

    # ===================== CONTENTS ====================================== #
    R.doc.add_page_break()
    R.doc.add_heading("Contents", level=1)
    R.toc_field()

    # ===================== 1. SCOPE & METHOD ============================= #
    R.h1("Scope and method")
    R.para("What the agent did, in order, and what each step guarantees.")

    R.h2("Pipeline")
    for i, (name, desc) in enumerate([
        ("Discover", "Read robots.txt, followed every Sitemap: directive, expanded "
                     "sitemap indexes, and supplemented with a breadth-first crawl "
                     "from the homepage so pages missing from the sitemap are still "
                     "found."),
        ("Fetch", "Polite asynchronous fetching that honours robots.txt Disallow "
                  "rules and Crawl-delay, follows redirects, retries rate-limits "
                  "and transient 5xx with backoff, and keeps only HTTP 200 HTML."),
        ("Canonicalise", "Normalised URLs, stripped tracking parameters, folded "
                         "rel=canonical duplicates together, recorded redirect "
                         "aliases so links to the pre-redirect URL still resolve."),
        ("Extract", "Stripped navigation, header, footer and sidebar regions, then "
                    "recorded each paragraph of body copy together with the exact "
                    "character spans occupied by links already inside it."),
        ("Graph", "Built the internal link graph, separating editorial (in-content) "
                  "links from site-wide navigation links, and ran PageRank over the "
                  "editorial graph only."),
        ("Vectorise", "TF-IDF over extracted body copy with title, H1 and URL slug "
                      "boosted, then a cosine similarity matrix. Fully "
                      "deterministic: the same crawl produces the same numbers."),
        ("Analyse", "Orphans, under-linked pages, keyword cannibalisation, "
                    "duplicate content, link saturation, broken links."),
        ("Recommend", "For each candidate source/target pair, located a verbatim "
                      "phrase already present in the source page's copy that also "
                      "distinguishes the target page from its siblings."),
    ], 1):
        R.rich([(f"{i}. {name} — ", {"bold": True}), (desc, {})],
               space_after=3, indent=0.1)

    R.h2("What is guaranteed, and what is not")
    R.para("These are the properties the tool enforces. They are the reason its "
           "orphan count is higher, and its recommendation count lower, than "
           "tools that count navigation links as editorial links.")
    for text in [
        "Every recommended target URL was fetched in this run and returned HTTP 200 "
        "as HTML. No target is a guess from a sitemap.",
        "Every anchor text in the “ready to implement” tier is a verbatim "
        "substring of the source page's own body copy, quoted with its surrounding "
        "sentence and character offset so it can be verified by hand.",
        "An anchor is never placed inside existing link text, never on a heading, "
        "and never in a paragraph that also appears on another page (that would be "
        "a template edit, producing a site-wide link rather than one editorial "
        "link).",
        "An anchor phrase must contain at least one token that distinguishes the "
        "target from its sibling pages, and the sentence hosting it must share "
        "vocabulary with the target. A phrase that describes many pages equally "
        "well is rejected even when it is present verbatim.",
        "Two pages that target the same primary keyword are never recommended to "
        "link to each other, and a contested keyword is never used as anchor text.",
        "Orphan status is computed from editorial inbound links only. Navigation "
        "and footer links are excluded, because counting them means nothing is ever "
        "an orphan. Both counts are shown side by side throughout.",
        "Nothing is fabricated. Any metric the crawl could not measure is absent "
        "from this document rather than estimated.",
    ]:
        R.bullet(text)

    R.h2("Limits of this run")
    lim = []
    if coverage.get("budget_exhausted"):
        lim.append(f"The page budget of {cfg['max_pages']} was reached, so the crawl "
                   f"is a sample of the site rather than all of it. Inbound links "
                   f"coming from pages that were not crawled are invisible to this "
                   f"run, which can make a page look like an orphan when it is not. "
                   f"Re-run with --max-pages "
                   f"{max(coverage.get('sitemap_urls', 0) + 100, cfg['max_pages'] * 2)} "
                   f"for a definitive picture.")
    if crawler.unfetched_discovered:
        lim.append(f"{len(crawler.unfetched_discovered)} URL(s) were discovered as "
                   f"link targets but never fetched, because the crawl budget ran "
                   f"out first. They are neither analysed nor reported as failures.")
    if empty_pages:
        lim.append(f"{len(empty_pages)} of {n} pages yielded almost no readable body "
                   f"text ({len(empty_pages) / n:.0%} of the crawl). Similarity, "
                   f"keyword and cannibalisation results for those specific pages "
                   f"are not reliable; they are listed in full later in this report.")
    if crawler.throttled:
        lim.append(f"{crawler.throttled} request(s) were rate-limited or returned a "
                   f"transient server error and had to be retried. A large number "
                   f"here means the server was under strain during the crawl.")
    if transient:
        lim.append(f"{len(transient)} URL(s) returned a rate-limit or 5xx status "
                   f"that survived every retry. They are listed separately from "
                   f"broken links, because they are most likely a temporary server "
                   f"condition rather than a defect in the site.")
    lim.append("Content rendered entirely by client-side JavaScript is not executed "
               "unless the run used --render. Any page whose copy is injected by "
               "script will appear thin.")
    lim.append("Similarity is lexical, not semantic. Two pages that cover the same "
               "subject in entirely different vocabulary will score as unrelated.")
    for text in lim:
        R.bullet(text)

    # ===================== 2. EXECUTIVE SUMMARY ========================== #
    R.h1("Executive summary")

    orphan_pct = (len(orphans) / n) if n else 0
    avg_ed = (total_ed / n) if n else 0

    R.h2("Headline findings")
    if orphan_pct >= 0.5:
        R.callout(
            "The site has no meaningful contextual linking layer",
            f"{len(orphans)} of {n} analysed pages ({orphan_pct:.0%}) have zero "
            f"in-content inbound links. Pages are reachable only through the "
            f"navigation menu, which means every page receives roughly the same "
            f"internal signal and none is distinguished as important. Across the "
            f"whole site there are {total_ed} editorial internal links, an average "
            f"of {avg_ed:.2f} per page. A healthy content site typically runs "
            f"three to ten.", kind="warn")
    elif orphan_pct >= 0.2:
        R.callout(
            "A substantial minority of pages have no contextual inbound links",
            f"{len(orphans)} of {n} analysed pages ({orphan_pct:.0%}) have zero "
            f"in-content inbound links, against {total_ed} editorial internal links "
            f"site-wide ({avg_ed:.2f} per page).", kind="warn")
    else:
        R.callout(
            "Internal linking is broadly healthy",
            f"{len(orphans)} of {n} analysed pages ({orphan_pct:.0%}) lack an "
            f"in-content inbound link, with {total_ed} editorial internal links "
            f"site-wide ({avg_ed:.2f} per page). The recommendations that follow "
            f"are refinements rather than repairs.")

    R.h2("The numbers")
    R.kv_table([
        ("Unique pages analysed (HTTP 200 HTML, after canonical merge)", f"{n}", ACCENT),
        ("URLs discovered in XML sitemaps", f"{len(crawler.sitemap_urls)}", None),
        ("Editorial (in-content) internal links", f"{total_ed}", None),
        ("Site-wide navigation / footer links (excluded from scoring)",
         f"{len(graph['boiler_edges'])}", None),
        ("Average editorial inbound links per page", f"{avg_ed:.2f}",
         BAD if avg_ed < 1 else (WARN if avg_ed < 3 else OK)),
        ("Orphan pages (0 editorial inbound links)", f"{len(orphans)}",
         BAD if orphans else OK),
        ("Under-linked pages (1-2 editorial inbound)", f"{len(underlinked)}",
         WARN if underlinked else OK),
        ("Duplicate-content URL pairs (near-identical copy)", f"{len(dupes)}",
         BAD if dupes else OK),
        ("Keyword cannibalisation pairs", f"{len(rivals)}", WARN if rivals else OK),
        ("Pages at or over the link-saturation ceiling", f"{len(saturated)}",
         WARN if saturated else OK),
        ("Pages carrying an excessive total link load", f"{len(link_heavy)}",
         WARN if link_heavy else OK),
        ("Broken internal links (error status, still linked)", f"{len(real_broken)}",
         BAD if real_broken else OK),
        ("Non-editorial pages excluded from analysis (archives, pagination, feeds)",
         f"{len(non_content)}", None),
        ("Pages with no readable body text", f"{len(empty_pages)}",
         WARN if empty_pages else OK),
        ("URLs that could not be fetched", f"{len(crawler.failures)}", None),
    ])

    R.h2("Recommendation tiers")
    R.para(f"{len(recs)} internal link recommendations were produced, in three tiers "
           f"of decreasing readiness:")
    R.table(
        ["Tier", "Count", "What it means", "Action"],
        [
            ["Ready to implement", str(len(high)),
             "A descriptive multi-word anchor is already present verbatim in the "
             "source page's own copy, and the sentence containing it is about the "
             "target page.",
             "Wrap the quoted words in a link. No copywriting needed."],
            ["Single-word anchor", str(len(single)),
             "Also verbatim, but one word only, so the surrounding context carries "
             "less of the meaning.",
             "Read the quoted sentence, then link if it fits."],
            ["Requires new copy", str(len(needs)),
             "The target is topically relevant, but the source page contains no "
             "phrase that both fits and distinguishes the target.",
             "A sentence must be written. Do not force the anchor in."],
        ],
        [1.5, 0.65, 3.0, 1.85],
        row_colors=[OK, WARN, MUTED],
    )

    R.h2("Recommended order of work")
    steps = []
    if dup_clusters:
        steps.append(("Resolve duplicate content first.",
                      f"{sum(len(c) for c in dup_clusters)} URLs across "
                      f"{len(dup_clusters)} cluster(s) serve substantially the same "
                      f"copy. Until one URL owns each topic, no amount of internal "
                      f"linking can make any of them rank, and their titles describe "
                      f"the wrong page so anchor text derived from them would be "
                      f"wrong too. These pages are excluded from the "
                      f"recommendations below for that reason."))
    if real_broken:
        steps.append(("Fix or remove broken internal links.",
                      f"{len(real_broken)} URL(s) return an error status and are "
                      f"still linked from live pages. Each one wastes crawl budget "
                      f"and sends visitors to an error page. This is the cheapest "
                      f"fix in the report."))
    if high:
        steps.append((f"Implement the {len(high)} ready-to-use links.",
                      "Each is a copy-paste change: the anchor words already exist "
                      "on the page. Work down the priority order in section 4."))
    if rivals:
        steps.append((f"Decide the winner for each of the {len(rivals)} competing "
                      f"page pairs.",
                      "Keep the stronger page as the target for the shared keyword "
                      "and either consolidate or re-angle the other. Until this is "
                      "settled, links pointing at either page split the signal."))
    if orphans:
        steps.append((f"Place links to the {len(orphans)} orphan pages.",
                      "Section 5 marks which orphans are already covered by a "
                      "recommendation and which still need manual placement."))
    if needs:
        steps.append((f"Commission copy for the {len(needs)} high-value gaps.",
                      "These are relevant targets with no suitable existing phrase "
                      "on the source page."))
    for i, (title, body) in enumerate(steps, 1):
        R.rich([(f"Step {i}. ", {"bold": True, "color": ACCENT}),
                (title + " ", {"bold": True}), (body, {})],
               space_after=5, indent=0.1)
    if not steps:
        R.para("No priority actions: the crawl found no duplicate content, no broken "
               "internal links, no orphans and no competing pages.")

    # ===================== 3. CRAWL RECORD =============================== #
    R.h1("Crawl record and data quality")
    R.para("Everything the crawl noticed about itself. Read this before acting on "
           "any number in the report: it states what was and was not measured.")

    R.h2("Coverage")
    cov = coverage
    if cov["sitemap_urls_not_analyzed"] and cov["budget_exhausted"]:
        R.callout("Coverage is partial — read before acting",
                  f"{cov['sitemap_urls_not_analyzed']} of {cov['sitemap_urls']} "
                  f"sitemap URLs were not analysed because the page budget was "
                  f"reached. Inbound links from those pages are invisible here, so "
                  f"some pages listed as orphans may not be orphans. Re-run with "
                  f"--max-pages {cov['sitemap_urls'] + 100} for a definitive list.",
                  kind="warn")
    elif cov["sitemap_urls_not_analyzed"]:
        R.callout("Coverage is complete for every analysable page",
                  f"All {n} pages returning HTTP 200 HTML were crawled and "
                  f"cross-linked. {cov['sitemap_urls_not_analyzed']} sitemap URL(s) "
                  f"were not analysed because they error, redirect elsewhere, or are "
                  f"not HTML — they cannot hold or receive a link. Orphan status "
                  f"below is definitive.")
    else:
        R.callout("Coverage is complete",
                  f"All {cov['sitemap_urls']} sitemap URLs were analysed, so orphan "
                  f"status in this report is definitive.")
    if cov.get("examples"):
        R.para("Sitemap URLs not analysed:", bold=True, size=10, space_after=2)
        for u in cov["examples"][:12]:
            R.bullet(_truncate(u, 110))

    if crawler.notes:
        R.h2("Crawl notes")
        R.para("Each of these is a decision the crawler made and is disclosing, not "
               "an error.", size=9.5, color=MUTED)
        for note in crawler.notes:
            R.bullet(note)

    R.h2("Content extraction")
    modes = {}
    for p in pages.values():
        modes[p.extraction_mode] = modes.get(p.extraction_mode, 0) + 1
    R.para("Body copy is isolated by stripping navigation and template regions. "
           "Where the primary method clearly misfired on this site's markup, a "
           "safer fallback was used and recorded, so no page is silently emptied.")
    R.table(["Extraction mode", "Pages", "Meaning"],
            [[m, str(c), {
                "normal": "Class-and-structure based extraction succeeded.",
                "structural-only": "Class-name heuristics removed too much, so only "
                                   "unambiguous structural chrome (nav/header/"
                                   "footer/aside) was stripped.",
                "raw-text": "The page uses almost no paragraph markup, so visible "
                            "text was taken wholesale and split into sentences.",
                "rendered": "Static HTML was thin; content came from a headless "
                            "browser render instead.",
             }.get(m, "—")] for m, c in sorted(modes.items())],
            [1.5, 0.7, 4.8])

    if reject_stats:
        R.h2("Precision filters")
        R.para("Candidates the tool refused to publish. A filter that reports "
               "nothing is indistinguishable from a filter that is not running, so "
               "the counts are stated here.")
        labels = {
            "ambiguous_phrases_dropped":
                "Anchor phrases dropped for describing more than one page equally "
                "well",
            "same_primary_keyword":
                "Pairs blocked: both pages target the same primary keyword",
            "no_distinctive_phrase":
                "Pairs skipped: no phrase distinguishes the target from its siblings",
            "phrase_is_contested_keyword":
                "Pairs skipped: the only anchor available was a contested keyword",
            "sentence_off_topic":
                "Verbatim matches rejected: the sentence is not about the target",
            "anchor_conflict":
                "Anchors rejected: the source already links those exact words "
                "elsewhere",
            "anchor_reused_on_source":
                "Anchors rejected to stop one page carrying the same anchor twice",
        }
        rows = [[labels.get(k, k), str(v)]
                for k, v in sorted(reject_stats.items(), key=lambda kv: -kv[1]) if v]
        if rows:
            R.table(["Filter", "Rejected"], rows, [5.4, 1.4],
                    aligns=[None, WD_ALIGN_PARAGRAPH.CENTER])

    # ===================== 4. RECOMMENDATIONS ============================ #
    R.h1("Link recommendations")

    sitewide_share = len(graph["sitewide"]) / n if n else 0
    if sitewide_share >= 0.5:
        R.callout(
            "Why this site yields few contextual recommendations",
            f"{len(graph['sitewide'])} of {n} crawled pages are already linked from "
            f"a site-wide menu or sidebar, so they are excluded as targets: there is "
            f"no value in recommending a link to a page that every page already "
            f"links to. This site's internal linking is navigation-driven rather "
            f"than contextual, and the meaningful finding is the orphan and "
            f"editorial-link picture, not the recommendation count. To analyse "
            f"contextual links regardless of navigation, re-run with "
            f"--boilerplate-ratio 0.95.", kind="warn")

    R.h2(f"Tier 1 — ready to implement ({len(high)})")
    if not high:
        R.para("No recommendation reached this tier: no source page contained a "
               "multi-word phrase that both appears verbatim and distinguishes a "
               "relevant target page from its siblings.", italic=True)
    else:
        R.para("Each entry below is a change that can be made without writing "
               "anything new. The quoted sentence is already on the source page; "
               "the words in brackets are the ones to wrap in a link. Character "
               "offsets are given so the exact occurrence can be located even when "
               "the same words appear more than once on the page.")
        SHOW = 60
        for r in high[:SHOW]:
            R.h3(f"#{r['priority']}  →  {_truncate(r['target_url'], 88)}")
            R.rich([("Source page: ", {"bold": True, "size": 9.5}),
                    (r["source_url"], {"size": 9.5})], space_after=1, indent=0.1)
            R.rich([("Target page: ", {"bold": True, "size": 9.5}),
                    (r["target_url"], {"size": 9.5})], space_after=1, indent=0.1)
            R.rich([("Anchor text: ", {"bold": True, "size": 9.5}),
                    (r["anchor_text"], {"size": 9.5, "mono": True, "bold": True,
                                        "color": ACCENT})],
                   space_after=1, indent=0.1)
            R.rich([("Location: ", {"bold": True, "size": 9.5}),
                    (f"paragraph block {r['block_index']}, characters "
                     f"{r['char_start']}–{r['char_end']}", {"size": 9.5})],
                   space_after=3, indent=0.1)
            if r.get("anchor_omits"):
                R.rich([("Worth extending: ", {"bold": True, "size": 9.5,
                                               "color": WARN}),
                        (f"this anchor does not mention "
                         f"“{r['anchor_omits']}”, which the target's URL is built "
                         f"around. The link points at the right page, but a reader "
                         f"clicking these words is not told that is where they are "
                         f"going. Extend the anchor if the sentence allows it.",
                         {"size": 9.5})],
                       space_after=3, indent=0.1)
            R.para("Existing sentence — already on the page, link the "
                   "bracketed words:", size=9.5, bold=True, space_after=2,
                   indent=0.1)
            ctx = r["context_sentence"]
            a = r["anchor_text"]
            pos = ctx.find(a)
            parts: list[tuple[str, dict]] = []
            if pos >= 0:
                if ctx[:pos]:
                    parts.append((ctx[:pos], {"size": 9.5, "italic": True}))
                parts.append(("[" + a + "]", {"size": 9.5, "bold": True,
                                              "color": ACCENT, "highlight": True}))
                if ctx[pos + len(a):]:
                    parts.append((ctx[pos + len(a):], {"size": 9.5, "italic": True}))
            else:
                parts.append((ctx, {"size": 9.5, "italic": True}))
            R.rich(parts, space_after=3, indent=0.3)
            R.rich([("Why: ", {"bold": True, "size": 9.5}),
                    (r["reason"], {"size": 9.5}),
                    (f"  —  score {r['score']:.3f}, similarity "
                     f"{r['similarity']:.3f}", {"size": 9.5, "color": MUTED})],
                   space_after=8, indent=0.1)
        if len(high) > SHOW:
            R.para(f"…and {len(high) - SHOW} more, in recommendations.xlsx "
                   f"(filter confidence = high).", italic=True, color=MUTED)

    if single:
        R.h2(f"Tier 2 — single-word anchors, verify context ({len(single)})")
        R.para("The anchor is genuinely present on the source page, but it is one "
               "word, so it can appear in a sentence that is not really about the "
               "target. Read the sentence, then skip any that do not fit.")
        R.table(["#", "Source", "Target", "Anchor", "Sentence it appears in"],
                [[str(r["priority"]), _truncate(r["source_url"], 46),
                  _truncate(r["target_url"], 46), r["anchor_text"],
                  _truncate(r["context_sentence"], 150)] for r in single[:45]],
                [0.4, 1.5, 1.5, 0.95, 2.65], size=8)
        if len(single) > 45:
            R.para(f"…and {len(single) - 45} more in recommendations.xlsx.",
                   italic=True, color=MUTED)

    if needs:
        R.h2(f"Tier 3 — relevant targets that need new copy ({len(needs)})")
        R.para("These targets are topically related to the source page, but the "
               "source contains no phrase that both reads naturally and identifies "
               "the target. The anchor shown is the target's own H1 or title: a "
               "sentence has to be written to host it. Do not force it into an "
               "existing sentence — that is how the anchor stops matching the "
               "content around it.")
        R.table(["#", "Source", "Target", "Suggested anchor", "Sim.", "Why"],
                [[str(r["priority"]), _truncate(r["source_url"], 42),
                  _truncate(r["target_url"], 42), _truncate(r["anchor_text"], 40),
                  f"{r['similarity']:.2f}", _truncate(r["reason"], 90)]
                 for r in needs[:45]],
                [0.4, 1.45, 1.45, 1.2, 0.45, 2.05], size=8)
        if len(needs) > 45:
            R.para(f"…and {len(needs) - 45} more in recommendations.xlsx.",
                   italic=True, color=MUTED)

    # ===================== 5. ORPHANS ==================================== #
    R.h1("Orphan pages")
    R.para("A page with zero editorial inbound links receives no internal signal "
           "from the site's own content. The navigation column shows whether it is "
           "reachable at all through site-wide chrome: a page with zero in both "
           "columns is discoverable only through the sitemap.")
    if orphans:
        fixed = {r["target_url"] for r in recs}
        rows, colors = [], []
        for o in orphans[:90]:
            if o["noindex"]:
                status, col = "n/a — noindex by design", MUTED
            elif o["url"] in dup_page_set:
                status, col = "n/a — duplicate, fix first", MUTED
            elif o["url"] in fixed:
                status, col = "Yes — link queued", OK
            else:
                status, col = "NO — needs manual placement", BAD
            rows.append([_truncate(o["url"], 62), "0",
                         str(o["inbound_boilerplate"]), str(o["word_count"]),
                         _truncate(o["title"] or "(no title)", 42), status])
            colors.append({5: col})
        R.table(["URL", "Editorial in", "Nav in", "Words", "Title", "Fix queued?"],
                rows, [2.25, 0.62, 0.5, 0.5, 1.55, 1.38], row_colors=colors, size=8)
        if len(orphans) > 90:
            R.para(f"…and {len(orphans) - 90} more in orphans.xlsx.",
                   italic=True, color=MUTED)
    else:
        R.para("None found. Every analysed content page has at least one in-content "
               "inbound link.", italic=True, color=OK)

    if underlinked:
        R.h2(f"Under-linked pages ({len(underlinked)})")
        R.para("One or two editorial inbound links. Not orphaned, but not "
               "reinforced either.")
        R.table(["URL", "Editorial in", "Words", "Internal PageRank"],
                [[_truncate(u["url"], 74), str(u["inbound_editorial"]),
                  str(u["word_count"]), f"{u['pagerank']:.5f}"]
                 for u in underlinked[:45]],
                [3.6, 0.85, 0.7, 1.65], size=8)
        if len(underlinked) > 45:
            R.para(f"…and {len(underlinked) - 45} more in orphans.xlsx "
                   f"(status = under-linked).", italic=True, color=MUTED)

    if non_content:
        R.h2(f"Pages excluded as non-editorial ({len(non_content)})")
        R.para("Paginated archives, tag and category listings, search results and "
               "feeds are real pages, but they are template output rather than "
               "content. They have no in-content inbound links by design and never "
               "will, so listing them as orphans would bury the pages that can "
               "actually be fixed. They are excluded from orphan counts, from "
               "cannibalisation, and from target candidacy — but they were still "
               "crawled, and links found on them still count.")
        R.table(["URL", "Kind", "Editorial in", "Nav in", "Words"],
                [[_truncate(r["url"], 70), r["kind"], str(r["inbound_editorial"]),
                  str(r["inbound_boilerplate"]), str(r["word_count"])]
                 for r in non_content[:45]],
                [3.35, 0.95, 0.85, 0.7, 0.95], size=8)
        if len(non_content) > 45:
            R.para(f"…and {len(non_content) - 45} more.", italic=True,
                   color=MUTED)

    # ===================== 6. DUPLICATES ================================= #
    if dup_clusters:
        # A cluster is only VERIFIED duplicate content if at least one of its
        # pairs was confirmed on pages the crawler could actually read. A
        # cluster built entirely from `needs-verification` pairs is, as far as
        # this run knows, a set of URLs that all returned the same unrendered
        # shell — telling someone to 301 those would destroy working pages.
        _crit_pages = {u for c in dupes for u in (c["page_a"], c["page_b"])}
        verified = [cl for cl in dup_clusters if any(u in _crit_pages for u in cl)]
        unverified = [cl for cl in dup_clusters if not any(u in _crit_pages for u in cl)]

        R.h1("Duplicate content")

        if verified:
            R.callout("Fix this before anything else in the report",
                      f"{sum(len(c) for c in verified)} URLs across "
                      f"{len(verified)} cluster(s) serve substantially the same copy "
                      f"— cosine similarity at or above "
                      f"{cfg['duplicate_similarity']:.2f} on extracted body text, from "
                      f"{len(dupes)} pairwise matches. While several URLs publish the "
                      f"same page, none can rank for the topic its own URL implies, and "
                      f"no amount of internal linking changes that.", kind="warn")
            R.para("These pages are excluded from the link recommendations, because "
                   "their titles and H1s describe the wrong content — so any anchor "
                   "text derived from them would be wrong too.")
            for i, cl in enumerate(verified, 1):
                R.h3(f"Cluster {i} — {len(cl)} URLs serving the same content")
                R.table(["URL", "Words", "Title"],
                        [[_truncate(u, 66),
                          str(pages[u].word_count if u in pages else "?"),
                          _truncate((pages[u].title if u in pages else "") or "(no title)", 50)]
                         for u in cl],
                        [3.2, 0.6, 3.0], size=8)
            R.para("Fix: decide which single URL should own each topic, give it unique "
                   "copy, and 301 or rel=canonical the rest to it. Then re-run this "
                   "audit — the orphan and recommendation numbers will change "
                   "substantially.", bold=True)

        if unverified:
            R.callout("Not confirmed — do not consolidate on this evidence",
                      f"{sum(len(c) for c in unverified)} URLs across "
                      f"{len(unverified)} cluster(s) returned near-identical text, but "
                      f"NONE of those pages could be parsed with the normal content "
                      f"extractor. That is the signature of a site whose content is "
                      f"rendered by JavaScript: the crawler read the same empty "
                      f"template from every URL and cannot tell them apart. This is "
                      f"probably not a duplicate-content problem at all.", kind="warn")
            for i, cl in enumerate(unverified, 1):
                R.h3(f"Unconfirmed cluster {i} — {len(cl)} URLs the crawler could not read")
                R.table(["URL", "Words", "Title"],
                        [[_truncate(u, 66),
                          str(pages[u].word_count if u in pages else "?"),
                          _truncate((pages[u].title if u in pages else "") or "(no title)", 50)]
                         for u in cl],
                        [3.2, 0.6, 3.0], size=8)
            R.para("Next step: re-run this audit with --render so the real content "
                   "is loaded, then check whether these URLs still look identical. "
                   "Only treat them as duplicates if they do.", bold=True)

    # ===================== 7. CANNIBALISATION ============================ #
    R.h1("Keyword cannibalisation and competing pages")
    R.para("These pairs target the same primary keyword or overlap heavily in "
           "content. They are excluded from the recommendations above so the tool "
           "never reinforces one competitor with the keyword the other is trying to "
           "win, and a contested keyword is never used as anchor text anywhere.")
    if rivals:
        rows, colors = [], []
        for c in rivals[:55]:
            rows.append([
                c["severity"].title(),
                _truncate(c["shared_keyword"], 34),
                f"{_truncate(c['page_a'], 44)}\n({c['words_a']} words, "
                f"{c['inbound_a']} inbound)",
                f"{_truncate(c['page_b'], 44)}\n({c['words_b']} words, "
                f"{c['inbound_b']} inbound)",
                f"{c['similarity']:.3f}",
            ])
            colors.append({0: _sev_color(c["severity"])})
        R.table(["Severity", "Shared keyword", "Page A", "Page B", "Similarity"],
                rows, [0.72, 1.35, 2.05, 2.05, 0.63], row_colors=colors, size=7.5)
        if len(rivals) > 55:
            R.para(f"…and {len(rivals) - 55} more in cannibalization.xlsx.",
                   italic=True, color=MUTED)
        R.para("Fix for each pair: keep the stronger page — more body copy, more "
               "editorial inbound links — as the single target for the shared "
               "keyword, and either 301 the weaker one to it or re-angle it onto a "
               "genuinely different keyword. Point every internal link about that "
               "keyword at the page you keep.", bold=True)
    else:
        R.para("No competing pages detected. No two analysed pages share a derived "
               "primary keyword with enough body-text overlap to indicate they are "
               "competing.", italic=True, color=OK)

    # ===================== 8. LINK LOAD ================================== #
    if saturated or link_heavy:
        R.h1("Excessive linking")
        if saturated:
            R.h2(f"In-content link saturation ({len(saturated)})")
            R.para("These pages already carry a heavy in-content link load, so no "
                   "new links were recommended from them. Beyond roughly one "
                   f"in-content link per {cfg['words_per_link']} words, each "
                   f"additional link dilutes the others and the copy starts reading "
                   f"as a link list.")
            R.table(["URL", "Editorial out", "Words", "Links per 100 words"],
                    [[_truncate(s["url"], 72), str(s["outbound_editorial"]),
                      str(s["word_count"]),
                      f"{(s['outbound_editorial'] / s['word_count'] * 100):.2f}"
                      if s["word_count"] else "—"]
                     for s in saturated[:35]],
                    [3.5, 0.95, 0.7, 1.65], size=8)
        if link_heavy:
            R.h2(f"Excessive total link load ({len(link_heavy)})")
            R.para("Total <a href> count on the page, navigation and footer "
                   "included. This is what a crawler actually sees, and the "
                   "editorial-only count above cannot detect it: a page can pass "
                   "the saturation test while carrying a several-hundred-link "
                   "mega-menu that dilutes every link on it.")
            R.table(["URL", "Total links", "Of which editorial", "Words"],
                    [[_truncate(r["url"], 72), str(r["links_total"]),
                      str(r["outbound_editorial"]), str(r["word_count"])]
                     for r in link_heavy[:35]],
                    [3.5, 0.95, 1.35, 0.7], size=8)

    # ===================== 9. BROKEN LINKS =============================== #
    if real_broken or malformed or transient:
        R.h1("Broken and malformed internal links")
    if real_broken:
        R.h2(f"Broken internal links ({len(real_broken)})")
        R.para("These URLs return an error status but are still linked from live "
               "pages on the site. Every one wastes crawl budget and sends visitors "
               "to an error page.")
        R.table(["Broken URL", "Status", "Linked from", "In sitemap?"],
                [[_truncate(b["url"], 74), b["status"],
                  f"{b['referring_pages']} page(s)",
                  "yes" if b["in_sitemap"] else "no"] for b in real_broken[:45]],
                [3.7, 0.8, 1.15, 1.15], size=8)
        if len(real_broken) > 45:
            R.para(f"…and {len(real_broken) - 45} more. Full referring-page "
                   f"lists are in broken_links.xlsx.", italic=True, color=MUTED)
        else:
            R.para("Full referring-page lists are in broken_links.xlsx.",
                   italic=True, color=MUTED)
    if transient:
        R.h2(f"Server errors during the crawl — not necessarily broken "
             f"({len(transient)})")
        R.para("These URLs returned a rate-limit or 5xx status that survived every "
               "retry. That usually means the server was unavailable or throttling "
               "during this crawl, not that the link is broken. They are listed "
               "apart so a temporary condition is not reported to you as a site "
               "defect — re-check them before acting.")
        R.table(["URL", "Status", "Linked from"],
                [[_truncate(b["url"], 82), b["status"],
                  f"{b['referring_pages']} page(s)"] for b in transient[:30]],
                [4.4, 0.9, 1.5], size=8)
    if malformed:
        R.h2(f"Malformed hrefs in the markup ({len(malformed)})")
        R.para("These use a single slash after the scheme "
               "(https:/example.com/page instead of https://example.com/page). "
               "Browsers and crawlers resolve them relative to the current page, "
               "producing URLs like /section/https:/example.com/page. Some servers "
               "answer HTTP 200 on those, which creates duplicate indexable URLs.")
        R.table(["Malformed href in page source", "Appears on"],
                [[_truncate(h, 92), f"{len(s)} page(s)"]
                 for h, s in list(malformed.items())[:25]],
                [5.4, 1.4], size=8)

    # ===================== 10. THIN PAGES ================================ #
    if empty_pages:
        R.h1("Pages with no readable content")
        R.para(f"{len(empty_pages)} page(s) returned HTTP 200 but yielded almost no "
               f"body text (under {cfg['min_content_words']} words after template "
               f"removal). Either they are genuinely thin, or their content is "
               f"rendered client-side by JavaScript, which this crawl did not "
               f"execute. Similarity, keyword and cannibalisation results for these "
               f"pages are not reliable, and none was used as a recommendation "
               f"source. Re-run with --render to resolve the JavaScript case.")
        R.table(["URL", "Words", "Extraction mode"],
                [[_truncate(u, 78),
                  str(pages[u].word_count if u in pages else "?"),
                  pages[u].extraction_mode if u in pages else "—"]
                 for u in empty_pages[:45]],
                [4.4, 0.7, 1.7], size=8)
        if len(empty_pages) > 45:
            R.para(f"…and {len(empty_pages) - 45} more; see crawl_data.json, "
                   f"word_count field.", italic=True, color=MUTED)

    # ===================== 11. AUTHORITY ================================= #
    R.h1("Internal authority distribution")
    R.para("PageRank computed over the editorial link graph only — navigation "
           "and footer links are excluded, because including them makes every page "
           "look equally important and tells you nothing. On a site with very few "
           "editorial links these values will be nearly uniform, which is itself "
           "the finding.")
    top = sorted(urls, key=lambda u: -pages[u].pagerank)[:20]
    R.table(["URL", "Internal PageRank", "Editorial in", "Editorial out", "Words"],
            [[_truncate(u, 62), f"{pages[u].pagerank:.5f}",
              str(pages[u].inbound_editorial), str(pages[u].outbound_editorial),
              str(pages[u].word_count)] for u in top],
            [2.85, 1.25, 0.85, 0.9, 0.95], size=8)

    if crawler.failures:
        R.h2(f"URLs that could not be used ({len(crawler.failures)})")
        R.table(["URL", "Reason"],
                [[_truncate(u, 88), str(why)]
                 for u, why in list(crawler.failures.items())[:45]],
                [5.0, 1.8], size=8)
        if len(crawler.failures) > 45:
            R.para(f"…and {len(crawler.failures) - 45} more in summary.json "
                   f"(fetch_failures).", italic=True, color=MUTED)

    # ===================== 12. GSC ======================================= #
    if gsc:
        R.h1("Search Console / analytics data")
        R.para(f"A --gsc-csv export was joined onto the crawl: "
               f"{gsc['matched']} row(s) matched a crawled page, "
               f"{gsc['unmatched']} row(s) had no matching crawled page. Matched "
               f"impressions were blended into the recommendation score, so pages "
               f"already earning impressions are prioritised as link targets.")
        if gsc.get("unmatched_samples"):
            R.para("Sample of unmatched URLs — check the export covers this same "
                   "host:", bold=True, size=10, space_after=2)
            for u in gsc["unmatched_samples"]:
                R.bullet(_truncate(u, 105))
        opp = [o for o in orphans if o.get("gsc_impressions", 0) > 0]
        if opp:
            R.h2("Highest-value linking gaps")
            R.para("Orphan pages that are already earning search impressions. These "
                   "are the best return on a single internal link: demand exists "
                   "and the site is giving the page no internal support.")
            R.table(["URL", "Impressions", "Words"],
                    [[_truncate(o["url"], 78), f"{o['gsc_impressions']:.0f}",
                      str(o["word_count"])] for o in opp[:20]],
                    [4.4, 1.2, 1.2], size=8)

    # ===================== 13. HOW TO VERIFY ============================= #
    R.h1("How to verify this report")
    R.para("Every claim here is checkable without trusting the tool. Three "
           "independent routes:")
    R.h2("By hand, in a browser")
    R.para("Take any Tier 1 recommendation. Open the source page, use Ctrl-F for "
           "the quoted sentence, and confirm the bracketed words are present and "
           "are not already a link. The character offsets identify which occurrence "
           "is meant when the same words appear more than once.")
    R.h2("Automatically, over the network")
    R.para("Run verify_report.py against the output folder. It re-fetches every "
           "source page live and re-checks each anchor, each sentence, and each "
           "structural cap using BeautifulSoup's own text extraction — "
           "deliberately not the agent's extraction code, so a bug in the agent "
           "cannot hide itself from the check.")
    R.h2("Against the raw data")
    R.para("Every table in this document has a machine-readable counterpart. Nothing "
           "shown here is computed in the document itself.")
    R.table(["File", "Contents"],
            [["recommendations.xlsx", "Every recommendation with anchor, sentence, "
                                     "block index and character offsets"],
             ["orphans.xlsx", "Orphan and under-linked pages, with both inbound "
                             "counts"],
             ["cannibalization.xlsx", "Competing and duplicate page pairs with "
                                     "evidence and severity"],
             ["broken_links.xlsx", "Error-status URLs, their referring pages, and "
                                  "broken-versus-transient classification"],
             ["non_editorial_pages.xlsx", "Archive, pagination, search and feed pages "
                                         "excluded from the analysis"],
             ["crawl_data.json", "Per-page extracted data: title, H1, word count, "
                                 "link counts, PageRank, keyword, top terms"],
             ["summary.json", "Machine-readable totals, full configuration, and "
                              "every crawl note"]],
            [1.9, 4.9], size=8.5)

    R.para(space_after=14)
    R.para(f"Report generated {now} by the Internal Linking Agent. "
           f"Analysis is deterministic: the same crawl of the same site produces "
           f"identical output.", size=9, color=MUTED, italic=True)

    R.save(path)
