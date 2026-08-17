#!/usr/bin/env python3
"""
seoaudit.py — Automated Technical SEO Audit (Semrush-style).

Crawls a website and reports issues the way a professional site-audit tool
(Semrush / Ahrefs) does: grouped into ERRORS / WARNINGS / NOTICES, counted
PER INSTANCE (per affected page / link / image / file), and scored with a
severity-weighted, frequency-based "Site Health" formula rather than a flat
per-category penalty.

Checks implemented
------------------
ERRORS
  * Broken internal & external links (per instance)
  * 4xx/5xx pages
  * Duplicate <title> tags
  * Duplicate meta descriptions
  * Duplicate content (near-duplicate body text)
  * Incorrect pages in sitemap.xml (redirect/broken/off-site)

WARNINGS
  * Unminified JavaScript / CSS files
  * Images without alt attributes
  * Low text-to-HTML ratio
  * Internal links with rel="nofollow"
  * Pages without an H1
  * Pages without a meta description
  * Missing <title>
  * Low word count
  * Title too long / too short
  * Meta description too long / too short
  * A link on an HTTPS page points to an HTTP URL
  * Mixed content (HTTP resources loaded on an HTTPS page)
  * Pages without a viewport meta tag
  * Pages canonicalised to a different URL (won't be indexed on their own)
  * www and non-www both serve content (duplicate host)

NOTICES
  * Resources formatted as a page link (linking to css/js/img/pdf via <a>)
  * Links with no anchor text
  * Non-descriptive anchor text ("click here", "read more", ...)
  * Permanent (301) redirects
  * Pages with more than one H1
  * Pages with more than one title tag
  * Pages without a charset declaration
  * Pages without a doctype
  * Pages with only one incoming internal link
  * Orphan pages (in sitemap but not internally linked)
  * Non-indexable pages (noindex / robots.txt)
  * robots.txt / sitemap notices
  * Missing / non-absolute / off-site canonical tags
  * Slow-loading pages
  * HSTS not enabled

The deliverable is a scheduled-audit DOCUMENT (Word .docx) with, for every
issue: Severity, Affected URL, Issue Type, and Recommended Action. A document
is always produced by default (auto-named seo-audit-<host>-<date>.docx) unless
--json is used; HTML is only written when --html is explicitly requested.

Install:
    pip install requests beautifulsoup4 python-docx
    (optional, for JS sites)  pip install playwright && playwright install chromium
    (python-docx is optional — without it the report is written as .rtf, which
     also opens in Word / Google Docs / LibreOffice.)

Usage:
    python main.py                                  # -> seo-audit-<host>-<date>.docx
    python main.py example.com
    python main.py https://example.com --max-pages 300
    python main.py example.com --doc audit.docx     # choose the document path
    python main.py example.com --html report.html   # also emit HTML (optional)
    python main.py example.com --json > audit.json  # data only, no document

Schedule it (Windows Task Scheduler), e.g. a weekly run:
    schtasks /create /tn "SEO Audit" /sc weekly /d MON /st 08:00 ^
      /tr "\"C:\\path\\.venv\\Scripts\\python.exe\" \"C:\\path\\main.py\" example.com"
"""

import sys
import io
import re
import json
import time
import gzip
import datetime
import html as html_lib
import argparse
import urllib.parse
import urllib.robotparser
import xml.etree.ElementTree as ET
from collections import Counter, defaultdict, deque
from concurrent.futures import ThreadPoolExecutor

# Ensure UTF-8 output on Windows consoles that default to cp1252
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except AttributeError:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

try:
    import requests
    requests.packages.urllib3.disable_warnings()
except ImportError:
    sys.exit("Install deps:  pip install requests beautifulsoup4")

try:
    from bs4 import BeautifulSoup
except ImportError:
    sys.exit("Install deps:  pip install requests beautifulsoup4")


UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
      "(KHTML, like Gecko) Chrome/124.0 Safari/537.36")

BROWSER_HEADERS = {
    "User-Agent": UA,
    "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
               "image/avif,image/webp,image/apng,*/*;q=0.8"),
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate",
    "Upgrade-Insecure-Requests": "1",
}

# A link is BROKEN only on these (definitively gone) statuses. Everything else
# (401/403/405/429/5xx/999/520-530 ...) is treated as bot-blocked / transient =
# "unverified", NEVER counted as broken.
BROKEN_STATUSES = {404, 410}

# Links with these extensions are never crawled as HTML (still link-checked).
NON_HTML_EXT = re.compile(
    r"\.(pdf|jpe?g|png|gif|webp|svg|ico|bmp|tiff?|mp[34]|m4[av]|mov|avi|wmv|"
    r"webm|zip|gz|tar|rar|7z|dmg|exe|msi|css|js|json|xml|rss|txt|csv|"
    r"docx?|xlsx?|pptx?|woff2?|ttf|eot)(\?|#|$)", re.I)

# Resource extensions that should never be the target of a normal page <a href>.
RESOURCE_EXT = re.compile(
    r"\.(css|js|json|pdf|jpe?g|png|gif|webp|svg|ico|bmp|tiff?|mp[34]|m4[av]|"
    r"mov|avi|wmv|webm|zip|gz|tar|rar|7z|dmg|exe|msi|docx?|xlsx?|pptx?|"
    r"woff2?|ttf|eot)(\?|#|$)", re.I)

NONDESC_ANCHORS = {
    "click here", "click", "here", "read more", "more", "learn more",
    "this", "link", "this link", "more info", "details", "continue",
    "continue reading", "go", "download",
}

# ---------------------------------------------------------------------------
# Severity tiers (Semrush terminology). `tier` drives the score weight; a check
# with 0 failures is displayed as PASSED regardless of its tier. The special
# "info" tier is NEVER scored — it is used only for the "could not verify"
# section (things we could not confirm either way).
# ---------------------------------------------------------------------------
TIER_WEIGHT = {"error": 5.0, "warning": 2.0, "notice": 1.0}   # "info" excluded on purpose
TIER_ORDER  = {"error": 0, "warning": 1, "notice": 2, "info": 3, "passed": 4}
TIER_COLOR  = {"error": "#dc2626", "warning": "#d97706", "notice": "#3b82f6",
               "info": "#64748b", "passed": "#16a34a"}
TIER_LABEL  = {"error": "ERROR", "warning": "WARNING", "notice": "NOTICE",
               "info": "UNVERIFIED", "passed": "PASSED"}
SECTION_TITLE = {"error": "ERRORS", "warning": "WARNINGS", "notice": "NOTICES",
                 "info": "COULD NOT VERIFY", "passed": "PASSED"}

# Thresholds (tunable via CLI where noted)
TITLE_MAX      = 60      # chars; longer = "title too long". Google truncates the
                         # SERP title around 580px ~= 60 chars, and the remediation
                         # advice below says 10-60, so the check must use 60 too.
                         # (Semrush's own cut-off is ~70, which silently passed
                         # 61-70 char titles that DO get truncated in the SERP.)
TITLE_MIN      = 10      # chars; shorter = "title too short"
DESC_MAX       = 160     # chars
DESC_MIN       = 70      # chars
LOW_WORDS      = 200     # visible words below this = "low word count"
TEXT_HTML_MIN  = 0.10    # text/HTML ratio below this = "low text-HTML ratio"

# Concavity of the per-check penalty. A check failing on fraction f of its units
# loses f**PENALTY_EXP of its weight (not f). 0.5 (square root) makes the FIRST
# occurrences of an issue count for more than later ones — reflecting that an
# issue appearing at all signals a site-wide problem. This is what makes the
# score behave like Semrush (concave in frequency) rather than a lenient linear
# ratio. 1.0 = pure linear; lower = harsher on any issue being present.
PENALTY_EXP    = 0.5


# ===========================================================================
# URL helpers
# ===========================================================================
def normalize_url(url):
    url = url.strip()
    if not re.match(r"^https?://", url, re.I):
        url = "https://" + url
    return url


def host_key(url):
    net = urllib.parse.urlparse(url).netloc.lower()
    if "@" in net:
        net = net.split("@")[-1]
    net = net.split(":")[0]
    if net.startswith("www."):
        net = net[4:]
    return net


def same_site(a, b):
    return host_key(a) == host_key(b)


def canon_url(url):
    p = urllib.parse.urlsplit(url)
    scheme = (p.scheme or "https").lower()
    netloc = p.netloc.lower()
    if netloc.startswith("www."):
        netloc = netloc[4:]
    if scheme == "http" and netloc.endswith(":80"):
        netloc = netloc[:-3]
    if scheme == "https" and netloc.endswith(":443"):
        netloc = netloc[:-4]
    path = p.path or "/"
    if len(path) > 1 and path.endswith("/"):
        path = path.rstrip("/")
    return urllib.parse.urlunsplit((scheme, netloc, path, p.query, ""))


def is_crawlable_html(url):
    return not NON_HTML_EXT.search(url)


def canonical_target(page, start_url):
    if not page.canonicals:
        return None
    href = page.canonicals[0]
    target = (href if re.match(r"^https?://", href, re.I)
              else urllib.parse.urljoin(page.url, href))
    return target if same_site(start_url, target) else None


def canonicalizes_away(page, start_url):
    t = canonical_target(page, start_url)
    return t is not None and canon_url(t) != canon_url(page.url)


ERROR_URL_RE = re.compile(
    r"/(404|not[-_]?found|page[-_]?not[-_]?found|error)(/|\.php|\.html?|$)", re.I)
ERROR_TITLE_RE = re.compile(r"^\s*(404\b|page not found|not found|error 404)", re.I)


def is_error_page(page):
    path = urllib.parse.urlsplit(page.url).path
    if ERROR_URL_RE.search(path):
        return True
    return bool(page.title and ERROR_TITLE_RE.search(page.title))


def is_thin(page):
    return (page.is_html and page.ok and not page.error
            and page.text_len < 200
            and not page.internal_links
            and not page.external_links
            and page.images_total == 0
            and len(page.h1s) == 0)


def thin_cause(page):
    if page.iframe_cross:
        return ("all content sits in a cross-origin iframe "
                f"({truncate(page.iframe_cross[0], 55)}) — search engines will not "
                "index it as part of this domain")
    if page.iframe_count:
        return "content is embedded in an iframe — poorly indexable"
    if page.spa_marker or page.script_count >= 3:
        return ("content is rendered client-side by JavaScript and is absent from "
                "the served HTML — run with --render to audit the rendered page")
    return "the page returned almost no HTML content"


def truncate(s, n=70):
    s = (s or "").strip()
    return s if len(s) <= n else s[: n - 1] + "…"


# ===========================================================================
# Page model + fetch/parse
# ===========================================================================
class Page:
    __slots__ = ("requested_url", "url", "status", "ok", "is_html", "elapsed",
                 "content_type", "error", "redirect_chain", "meta_refresh",
                 "title", "title_count", "meta_desc", "meta_desc_count",
                 "h1s", "canonicals", "robots_meta", "x_robots", "hsts",
                 "images_total", "images_missing_alt", "missing_alt_samples",
                 "internal_links", "external_links", "nofollow_internal",
                 "link_count", "empty_anchor", "empty_anchor_urls", "nondesc_anchor",
                 "resource_link_count", "http_from_https",
                 "assets", "text_len", "html_len", "word_count",
                 "content_sig", "script_count", "iframe_count", "spa_marker",
                 "iframe_cross", "rendered", "mixed_content", "has_viewport",
                 "has_charset", "has_doctype")

    def __init__(self, requested_url):
        self.requested_url = requested_url
        self.url = requested_url
        self.status = None
        self.ok = False
        self.is_html = False
        self.elapsed = 0.0
        self.content_type = ""
        self.error = None
        self.redirect_chain = []      # [(status, url), ...] hops before final
        self.meta_refresh = None
        self.title = None
        self.title_count = 0
        self.meta_desc = None
        self.meta_desc_count = 0
        self.h1s = []
        self.canonicals = []
        self.robots_meta = ""
        self.x_robots = ""
        self.hsts = None              # Strict-Transport-Security header value
        self.images_total = 0
        self.images_missing_alt = 0
        self.missing_alt_samples = []
        self.internal_links = set()
        self.external_links = set()
        self.nofollow_internal = set()
        self.link_count = 0           # raw count of valid <a href>
        self.empty_anchor = 0         # <a> with no anchor text
        self.empty_anchor_urls = []   # their targets, for unique-pattern counting
        self.nondesc_anchor = 0       # <a> with non-descriptive text
        self.resource_link_count = 0  # <a href> pointing at a resource file
        self.http_from_https = 0      # http:// links on an https page
        self.assets = set()           # absolute css/js URLs referenced
        self.text_len = 0             # visible body text length (chars)
        self.html_len = 0             # raw HTML length (chars)
        self.word_count = 0           # visible body word count
        self.content_sig = None       # token signature for near-dup detection
        self.script_count = 0
        self.iframe_count = 0
        self.spa_marker = False
        self.iframe_cross = []
        self.rendered = False
        self.mixed_content = 0        # http:// resources referenced on an https page
        self.has_viewport = False
        self.has_charset = False
        self.has_doctype = False


def fetch_page(session, url):
    page = Page(url)
    t0 = time.time()
    # Transient failures (timeout / connection reset — often just rate-limiting)
    # get one retry so a throttled response doesn't silently drop the page from
    # the crawl and make issue counts wobble between runs.
    last_transient = None
    for attempt in range(2):
        try:
            resp = session.get(url, timeout=20, allow_redirects=True, verify=False)
            page.elapsed = round(time.time() - t0, 3)
            page.status = resp.status_code
            page.url = resp.url
            page.ok = resp.status_code == 200
            page.content_type = resp.headers.get("Content-Type", "")
            page.x_robots = resp.headers.get("X-Robots-Tag", "")
            page.hsts = resp.headers.get("Strict-Transport-Security")
            page.redirect_chain = [(h.status_code, h.url) for h in resp.history]
            page.is_html = ("html" in page.content_type.lower()
                            or (not page.content_type and page.ok))
            if page.is_html and resp.text:
                _parse_html(page, resp.text)
            return page
        except (requests.exceptions.Timeout,
                requests.exceptions.ConnectionError) as e:
            last_transient = e
            time.sleep(0.6)                     # brief back-off, then retry once
            continue
        except requests.exceptions.TooManyRedirects:
            page.elapsed = round(time.time() - t0, 3)
            page.error = "Redirect loop (too many redirects)"
            return page
        except requests.exceptions.SSLError:
            page.elapsed = round(time.time() - t0, 3)
            page.error = "SSL error"
            return page
        except Exception as e:
            page.elapsed = round(time.time() - t0, 3)
            page.error = f"{type(e).__name__}: {str(e)[:80]}"
            return page
    page.elapsed = round(time.time() - t0, 3)
    page.error = ("Timeout (>20s)"
                  if isinstance(last_transient, requests.exceptions.Timeout)
                  else "Connection error / DNS failure")
    return page


_WORD_RE = re.compile(r"[^\W\d_]+", re.UNICODE)


def _content_signature(text):
    """A cheap near-duplicate fingerprint: the set of overlapping 5-word
    shingles of the normalised visible text. Two pages are near-duplicates when
    their shingle sets overlap heavily (Jaccard similarity)."""
    words = _WORD_RE.findall(text.lower())
    if len(words) < 8:
        return frozenset()
    shingles = {" ".join(words[i:i + 5]) for i in range(len(words) - 4)}
    return frozenset(shingles)


def _parse_html(page, text):
    page.html_len = len(text)
    page.has_doctype = text.lstrip()[:15].lower().startswith("<!doctype")
    soup = BeautifulSoup(text, "html.parser")

    body = soup.body
    visible = body.get_text(" ", strip=True) if body else ""
    page.text_len = len(visible)
    page.word_count = len(visible.split())
    page.content_sig = _content_signature(visible)
    page.script_count = len(soup.find_all("script"))

    iframes = soup.find_all("iframe")
    page.iframe_count = len(iframes)
    for fr in iframes:
        src = (fr.get("src") or "").strip()
        if src:
            abssrc = urllib.parse.urljoin(page.url, src)
            if abssrc.startswith(("http://", "https://")) and not same_site(page.url, abssrc):
                page.iframe_cross.append(abssrc)
    page.spa_marker = bool(re.search(
        r"""id=["'](root|app|__next|__nuxt)["']|__NEXT_DATA__|data-reactroot|"""
        r"""window\.__NUXT__|data-server-rendered|ng-version""", text, re.I))

    titles = soup.find_all("title")
    page.title_count = len(titles)
    if titles:
        page.title = titles[0].get_text(strip=True)

    descs = [m for m in soup.find_all("meta")
             if (m.get("name") or "").strip().lower() == "description"]
    page.meta_desc_count = len(descs)
    if descs:
        page.meta_desc = (descs[0].get("content") or "").strip()

    page.h1s = [h.get_text(strip=True) for h in soup.find_all("h1")]

    for link in soup.find_all("link"):
        rel = link.get("rel") or []
        if isinstance(rel, str):
            rel = rel.split()
        rel_l = [r.lower() for r in rel]
        href = (link.get("href") or "").strip()
        if any(r == "canonical" for r in rel_l) and href:
            page.canonicals.append(href)
        if "stylesheet" in rel_l and href:
            page.assets.add(urllib.parse.urljoin(page.url, href).split("#")[0])

    for s in soup.find_all("script", src=True):
        src = (s.get("src") or "").strip()
        if src:
            page.assets.add(urllib.parse.urljoin(page.url, src).split("#")[0])

    robots_vals = []
    for m in soup.find_all("meta"):
        name = (m.get("name") or "").strip().lower()
        if name in ("robots", "googlebot"):
            robots_vals.append((m.get("content") or "").lower())
        if name == "viewport":
            page.has_viewport = True
        if m.get("charset") or (m.get("http-equiv") or "").lower() == "content-type":
            page.has_charset = True
    page.robots_meta = ", ".join(v for v in robots_vals if v)

    refresh = soup.find("meta", attrs={"http-equiv": re.compile(r"^refresh$", re.I)})
    if refresh and refresh.get("content"):
        m = re.search(r"url\s*=\s*(.+)$", refresh["content"], re.I)
        if m:
            page.meta_refresh = urllib.parse.urljoin(page.url, m.group(1).strip().strip("'\""))

    # Images — exclude <noscript> lazy-load duplicates. A missing alt means the
    # attribute is ENTIRELY absent (alt="" is a valid decorative marker).
    imgs = [img for img in soup.find_all("img") if img.find_parent("noscript") is None]
    page.images_total = len(imgs)
    for img in imgs:
        if "alt" in img.attrs:
            continue
        src = (img.get("src") or img.get("data-src") or "").strip()
        if src.startswith("data:"):
            continue
        if img.get("width") in ("0", "1") or img.get("height") in ("0", "1"):
            continue
        page.images_missing_alt += 1
        if len(page.missing_alt_samples) < 5:
            page.missing_alt_samples.append(
                urllib.parse.urljoin(page.url, src or "(inline image)"))

    page_is_https = page.url.lower().startswith("https://")

    # Mixed content — an https page that loads a resource over plain http.
    if page_is_https:
        for tag, attr in (("img", "src"), ("script", "src"), ("link", "href"),
                          ("iframe", "src"), ("source", "src"), ("video", "src"),
                          ("audio", "src"), ("embed", "src")):
            for el in soup.find_all(tag):
                if (el.get(attr) or "").strip().lower().startswith("http://"):
                    page.mixed_content += 1

    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").strip()
        hl = href.lower()
        # Skip non-navigational links and MALFORMED contact links. A raw space in
        # an href, or a "tel"/"mailto"/"sms" prefix followed by any non-letter
        # (e.g. the typo `href="tel+1 (332) …"` — missing the colon), means it is
        # a phone/e-mail link, NOT a page. `[:+.\s]` after the scheme keeps real
        # pages like /telephone-services (which continue with a letter) crawlable.
        if (not href or " " in href
                or hl.startswith(("#", "javascript:", "data:", "vbscript:"))
                or re.match(r"(?:mailto|tel|sms|fax|callto|whatsapp|skype|viber)[:+.\s]", hl)):
            continue
        full = urllib.parse.urljoin(page.url, href)
        p = urllib.parse.urlsplit(full)
        if p.scheme not in ("http", "https"):
            continue
        full = urllib.parse.urlunsplit((p.scheme, p.netloc, p.path, p.query, ""))
        page.link_count += 1

        rel = a.get("rel") or []
        if isinstance(rel, str):
            rel = rel.split()
        rel_l = [r.lower() for r in rel]

        # Anchor-text quality. A link "has no anchor text" when it has NO visible
        # text, NO accessible name, and NO visual child carrying its own text
        # alternative.
        #
        # An image/icon link counts as NAMED only if the visual actually supplies
        # a text alternative — <img alt="…"> or <svg><title>…</title>. Excusing
        # EVERY link with a visual child (the previous rule) was wrong: an <img>
        # with a missing or empty alt, or a bare <svg>/<i> icon, gives Google
        # nothing to read, which is exactly the case this check exists to catch.
        # That blanket skip under-counted badly — 79 found vs Semrush's 252 on
        # americanwebbuilders.com, with all 79 landing on a single page.
        atext = a.get_text(" ", strip=True)
        aria = (a.get("aria-label") or a.get("title") or "").strip()
        if not atext and not aria:
            for el in a.find_all(True):
                if (el.get("aria-label") or el.get("title") or "").strip():
                    aria = "x"
                    break
                cls = " ".join(el.get("class") or [])
                if re.search(r"sr-only|visually-hidden|screen-reader", cls, re.I):
                    aria = "x"
                    break
        named_visual = False
        if not atext and not aria:
            for el in a.find_all(["img", "svg"]):
                if el.name == "img":
                    if (el.get("alt") or "").strip():
                        named_visual = True
                        break
                else:
                    st = el.find("title")
                    if st is not None and st.get_text(strip=True):
                        named_visual = True
                        break
        if not atext and not aria and not named_visual:
            page.empty_anchor += 1
            page.empty_anchor_urls.append(full)
        elif atext and atext.lower() in NONDESC_ANCHORS:
            page.nondesc_anchor += 1

        if RESOURCE_EXT.search(full):
            page.resource_link_count += 1
        if page_is_https and p.scheme == "http":
            page.http_from_https += 1

        if same_site(page.url, full):
            page.internal_links.add(full)
            if "nofollow" in rel_l:
                page.nofollow_internal.add(full)
        else:
            page.external_links.add(full)


# ===========================================================================
# Crawler
# ===========================================================================
def crawl(session, start_url, max_pages, workers, delay, log=True):
    seed = canon_url(start_url)
    seen = {seed}
    pages = {}
    link_sources = defaultdict(set)
    # Internal links keyed AS WRITTEN (www/non-www, http/https, trailing slash
    # preserved). link_sources folds those variants together via canon_url(),
    # which is right for most checks but hides redirects — links to the non-www
    # homepage looked identical to links to the canonical www one.
    raw_link_sources = defaultdict(set)
    frontier = deque([start_url])

    with ThreadPoolExecutor(max_workers=workers) as ex:
        while frontier and len(pages) < max_pages:
            wave = []
            while (frontier and len(wave) < workers * 3
                   and len(pages) + len(wave) < max_pages):
                wave.append(frontier.popleft())

            for page in ex.map(lambda u: fetch_page(session, u), wave):
                pages[canon_url(page.requested_url)] = page
                for link in sorted(page.internal_links):
                    link_sources[canon_url(link)].add(page.url)
                    raw_link_sources[link].add(page.url)
                    c = canon_url(link)
                    if c not in seen and is_crawlable_html(link):
                        seen.add(c)
                        frontier.append(link)
                for link in sorted(page.external_links):
                    link_sources[canon_url(link)].add(page.url)

            if log:
                sys.stderr.write(f"\r  crawled {len(pages)} pages, "
                                 f"{len(frontier)} queued ...   ")
                sys.stderr.flush()
            if delay:
                time.sleep(delay)

    if log:
        sys.stderr.write("\r" + " " * 60 + "\r")
        sys.stderr.flush()
    complete = len(frontier) == 0
    return pages, link_sources, raw_link_sources, complete


# ===========================================================================
# Headless-browser (Playwright) rendering
# ===========================================================================
def playwright_available():
    try:
        import playwright  # noqa: F401
        return True
    except ImportError:
        return False


def _render_fetch(ctx, url, wait_ms, start_url):
    page = Page(url)
    page.rendered = True
    pw_page = ctx.new_page()
    t0 = time.time()
    try:
        resp = pw_page.goto(url, wait_until="load", timeout=30000)
        page.elapsed = round(time.time() - t0, 3)
        try:
            pw_page.wait_for_load_state("networkidle", timeout=6000)
        except Exception:
            pass
        pw_page.wait_for_timeout(wait_ms)
        page.url = pw_page.url
        page.status = resp.status if resp else None
        page.ok = page.status == 200
        ct = ""
        if resp:
            try:
                ct = resp.headers.get("content-type", "")
                page.hsts = resp.headers.get("strict-transport-security")
            except Exception:
                pass
        page.content_type = ct
        page.is_html = "html" in ct.lower() or (not ct and bool(page.ok))
        if resp:
            chain, rf = [], resp.request.redirected_from
            while rf is not None:
                r2 = rf.response()
                chain.append(((r2.status if r2 else None), rf.url))
                rf = rf.redirected_from
            page.redirect_chain = list(reversed(chain))
        htmls, cross = [pw_page.content()], []
        for fr in pw_page.frames[1:]:
            try:
                if same_site(page.url, fr.url):
                    htmls.append(fr.content())
                elif fr.url.startswith(("http://", "https://")):
                    cross.append(fr.url)
            except Exception:
                pass
        if page.is_html:
            _parse_html(page, "\n".join(htmls))
        page.iframe_cross = list(dict.fromkeys(page.iframe_cross + cross))
    except Exception as e:
        page.elapsed = round(time.time() - t0, 3)
        page.error = f"render error: {str(e)[:80]}"
    finally:
        try:
            pw_page.close()
        except Exception:
            pass
    return page


def render_crawl(start_url, max_pages, delay, wait_ms, log=True):
    from playwright.sync_api import sync_playwright
    seed = canon_url(start_url)
    seen = {seed}
    pages, link_sources = {}, defaultdict(set)
    raw_link_sources = defaultdict(set)      # see crawl() — links as written
    frontier = deque([start_url])
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        ctx = browser.new_context(user_agent=UA, ignore_https_errors=True,
                                  viewport={"width": 1366, "height": 900})
        while frontier and len(pages) < max_pages:
            url = frontier.popleft()
            page = _render_fetch(ctx, url, wait_ms, start_url)
            pages[canon_url(page.requested_url)] = page
            for link in sorted(page.internal_links):
                link_sources[canon_url(link)].add(page.url)
                raw_link_sources[link].add(page.url)
                c = canon_url(link)
                if c not in seen and is_crawlable_html(link):
                    seen.add(c)
                    frontier.append(link)
            for link in sorted(page.external_links):
                link_sources[canon_url(link)].add(page.url)
            if log:
                sys.stderr.write(f"\r  rendered {len(pages)} pages, "
                                 f"{len(frontier)} queued ...   ")
                sys.stderr.flush()
            if delay:
                time.sleep(delay)
        browser.close()
    if log:
        sys.stderr.write("\r" + " " * 60 + "\r")
        sys.stderr.flush()
    return pages, link_sources, raw_link_sources, len(frontier) == 0


# ===========================================================================
# robots.txt + sitemap
# ===========================================================================
def fetch_robots(session, base_url):
    out = {"exists": False, "status": None, "text": "", "sitemaps": [],
           "parser": None, "issues": [], "blocks_all": False}
    robots_url = urllib.parse.urljoin(base_url, "/robots.txt")
    try:
        r = session.get(robots_url, timeout=12, verify=False)
        out["status"] = r.status_code
        if r.status_code == 200 and r.text.strip():
            out["exists"] = True
            out["text"] = r.text
            rp = urllib.robotparser.RobotFileParser()
            rp.parse(r.text.splitlines())
            out["parser"] = rp
            for line in r.text.splitlines():
                if line.lower().startswith("sitemap:"):
                    out["sitemaps"].append(line.split(":", 1)[1].strip())
            ua_all = re.search(r"user-agent:\s*\*(.*?)(?:\nuser-agent:|\Z)",
                               r.text, re.I | re.S)
            if ua_all and re.search(r"^\s*disallow:\s*/\s*$",
                                    ua_all.group(1), re.I | re.M):
                out["blocks_all"] = True
    except requests.RequestException as e:
        out["issues"].append(f"robots.txt fetch failed: {type(e).__name__}")
    return out


def fetch_sitemaps(session, sitemap_urls, base_url):
    out = {"found": [], "urls": [], "lastmods": {}, "issues": [], "count": 0}
    if sitemap_urls:
        queue, guessed = list(sitemap_urls), set()
    else:
        guessed = {urllib.parse.urljoin(base_url, p) for p in
                   ("/sitemap.xml", "/sitemap_index.xml", "/sitemap-index.xml")}
        queue = list(guessed)
    tried, seen_sm = [], set()

    def note(sm_url, msg):
        if sm_url not in guessed:
            out["issues"].append(f"{sm_url} — {msg}")

    while queue and len(out["found"]) < 25:
        sm_url = queue.pop(0)
        if sm_url in seen_sm:
            continue
        seen_sm.add(sm_url)
        tried.append(sm_url)
        try:
            r = session.get(sm_url, timeout=12, verify=False, headers=BROWSER_HEADERS)
        except requests.RequestException as e:
            note(sm_url, f"fetch failed ({type(e).__name__})")
            continue
        if r.status_code != 200:
            note(sm_url, f"HTTP {r.status_code}")
            continue
        content = r.content
        ctype = r.headers.get("Content-Type", "").lower()
        if sm_url.lower().endswith(".gz") or "gzip" in ctype:
            try:
                content = gzip.decompress(content)
            except Exception:
                pass
        head = content[:512].lstrip().lower()
        if head.startswith(b"<!doctype html") or head.startswith(b"<html"):
            note(sm_url, "returns an HTML page, not an XML sitemap")
            continue
        try:
            root = ET.fromstring(content)
        except ET.ParseError:
            note(sm_url, "not valid XML")
            continue
        out["found"].append(sm_url)
        tag = root.tag.lower()
        if tag.endswith("sitemapindex"):
            for sm in root.iter():
                if sm.tag.lower().endswith("loc") and sm.text:
                    queue.append(sm.text.strip())
        else:
            for url_el in root.iter():
                if not url_el.tag.lower().endswith("url"):
                    continue
                loc = lastmod = None
                for child in url_el:
                    t = child.tag.lower()
                    if t.endswith("loc") and child.text:
                        loc = child.text.strip()
                    elif t.endswith("lastmod") and child.text:
                        lastmod = child.text.strip()
                if loc:
                    out["urls"].append(loc)
                    if lastmod:
                        out["lastmods"][canon_url(loc)] = lastmod

    out["tried"] = tried
    out["count"] = len(out["urls"])
    return out


# ===========================================================================
# Link checking (external + uncrawled internal)
# ===========================================================================
def check_link(session, url):
    """Verify one URL. Returns (url, status, err, first_hop_status, redirected).
    first_hop_status/redirected let the caller detect permanent (301) redirects
    without a second request."""
    last_err = None
    for _ in range(2):
        try:
            r = session.get(url, timeout=15, allow_redirects=True,
                            verify=False, stream=True, headers=BROWSER_HEADERS)
            status = r.status_code
            first_hop = r.history[0].status_code if r.history else status
            redirected = bool(r.history)
            r.close()
            return url, status, None, first_hop, redirected
        except requests.exceptions.ConnectionError:
            last_err = "Connection error"
            continue
        except requests.exceptions.Timeout:
            return url, None, "Timeout", None, False
        except requests.exceptions.TooManyRedirects:
            return url, None, "Redirect loop", None, False
        except requests.exceptions.SSLError:
            return url, None, "SSL error", None, False
        except Exception as e:
            return url, None, type(e).__name__, None, False
    return url, None, last_err, None, False


def check_links(session, urls, workers):
    results = {}
    if not urls:
        return results
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for url, status, err, first_hop, redirected in ex.map(
                lambda u: check_link(session, u), urls):
            results[url] = (status, err, first_hop, redirected)
    return results


def link_verdict(status, err):
    if status in BROKEN_STATUSES:
        return "broken", f"HTTP {status}"
    if err and ("connection" in err.lower() or "dns" in err.lower()):
        return "broken", "connection refused / DNS failure"
    if status is not None and 200 <= status < 400:
        return "ok", None
    if status is not None:
        return "unverified", f"HTTP {status} (bot-blocked or transient)"
    return "unverified", err or "not verifiable"


def verify_reason(status, err):
    """Plain-language reason an unverifiable link could not be confirmed —
    shown in the 'could not verify' section so the report explains WHY."""
    if status is not None:
        if status in (401, 403, 405, 406, 999):
            return f"blocked to automated crawlers (HTTP {status}) — likely works in a browser"
        if status == 429:
            return "rate-limited by the server (HTTP 429)"
        if 500 <= status < 600:
            return f"server error / temporarily unavailable (HTTP {status})"
        return f"unexpected response (HTTP {status})"
    e = (err or "").lower()
    if "timeout" in e:
        return "request timed out (server too slow or blocking crawlers)"
    if "ssl" in e:
        return "SSL / certificate error"
    if "redirect" in e:
        return "redirect loop"
    if "connection" in e:
        return "connection error"
    return err or "could not connect"


# ===========================================================================
# Resource minification check
# ===========================================================================
def is_unminified(url, text):
    """Heuristic: minified files pack code onto very few, very long lines.
    Unminified files have many newlines and short average line length."""
    low = url.lower()
    if ".min." in low or "-min." in low:
        return False
    if not text or len(text) < 800:      # tiny files aren't worth flagging
        return False
    lines = text.count("\n") + 1
    if lines <= 8:
        return False                     # already on a handful of long lines
    avg_line = len(text) / lines
    # lots of indentation / comment markers is another strong tell
    indented = text.count("\n  ") + text.count("\n\t")
    return avg_line < 200 and (lines > 15 or indented > 10)


def check_resource(session, url):
    try:
        r = session.get(url, timeout=12, verify=False, headers=BROWSER_HEADERS,
                        stream=True)
        if r.status_code != 200:
            r.close()
            return url, None
        raw = r.raw.read(300000, decode_content=True) or b""
        r.close()
        text = raw.decode("utf-8", "replace")
        return url, is_unminified(url, text)
    except Exception:
        return url, None


def check_resources(session, urls, workers):
    results = {}
    if not urls:
        return results
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for url, unmin in ex.map(lambda u: check_resource(session, u), urls):
            if unmin is not None:
                results[url] = unmin
    return results


def check_host_variants(session, start_url):
    """Detect the www vs non-www duplicate-host problem: if BOTH hosts serve
    HTTP 200 without one redirecting to the other, every URL exists twice
    (this is exactly why crawlers like Semrush report ~2x the page/issue count
    on such sites). Returns {'duplicate': bool, 'www': url, 'nonwww': url}."""
    base = host_key(start_url)               # registrable host, www stripped
    www_url, nonwww_url = f"https://www.{base}/", f"https://{base}/"

    def probe(u):
        try:
            r = session.get(u, timeout=15, allow_redirects=True, verify=False)
            fh = urllib.parse.urlsplit(r.url).netloc.lower().split(":")[0]
            return r.status_code, fh
        except requests.RequestException:
            return None, None

    ws, wh = probe(www_url)
    ns, nh = probe(nonwww_url)
    duplicate = (ws == 200 and ns == 200 and wh and nh
                 and wh.startswith("www.") and not nh.startswith("www."))
    return {"duplicate": duplicate, "www": www_url, "nonwww": nonwww_url}


# ===========================================================================
# Duplicate-content detection (near-duplicate body text)
# ===========================================================================
def find_duplicate_content(pages, threshold=0.85):
    """Cluster pages whose visible-text shingle sets overlap >= threshold
    (Jaccard). Returns a list of clusters (each a list of page URLs)."""
    sigs = [(p, p.content_sig) for p in pages
            if p.content_sig and len(p.content_sig) >= 10]
    clusters, used = [], set()
    for i in range(len(sigs)):
        pi, si = sigs[i]
        if id(pi) in used:
            continue
        group = [pi]
        for j in range(i + 1, len(sigs)):
            pj, sj = sigs[j]
            if id(pj) in used:
                continue
            inter = len(si & sj)
            if not inter:
                continue
            union = len(si | sj)
            if union and inter / union >= threshold:
                group.append(pj)
                used.add(id(pj))
        if len(group) > 1:
            used.add(id(pi))
            clusters.append([p.url for p in group])
    return clusters


# ===========================================================================
# Analysis — one finding per check, counted per instance
# ===========================================================================
def finding(fid, name, tier, summary, items=None, failed=0, total=0, unit="pages"):
    display = tier if failed > 0 else "passed"
    return {"id": fid, "name": name, "tier": tier, "display": display,
            "summary": summary, "items": items or [],
            "failed": failed, "total": total, "unit": unit}


def analyze(start_url, pages, link_sources, robots, sitemap,
            link_status, resource_status, slow_threshold, external_checked,
            crawl_complete, max_pages, host_variants=None,
            raw_link_sources=None):
    findings = []
    seed = canon_url(start_url)

    # De-duplicate on FINAL (post-redirect) URL. Prefer the copy fetched directly.
    seen_final = {}
    for p in pages.values():
        if not (p.is_html and p.ok and not p.error):
            continue
        if is_error_page(p):
            continue
        key = canon_url(p.url)
        if key not in seen_final or not p.redirect_chain:
            seen_final[key] = p
    html_pages = list(seen_final.values())
    # On-page issues (missing meta/H1/alt/title, word count, text ratio, viewport,
    # …) are audited on EVERY indexable crawled page — exactly like Semrush. A
    # page that canonicalises to another URL still has these problems, and Semrush
    # reports them, so it must NOT be excluded here. (Excluding them was a bug:
    # a site that canonicalises every page to the homepage would show a perfect
    # score while actually being broken.)
    onpage = html_pages
    n_pages = max(1, len(onpage))

    # For DUPLICATE detection only, drop pages that defer indexing to a DIFFERENT
    # URL via rel=canonical (e.g. ?utm= variants → clean URL): the canonical
    # legitimately resolves the duplication, so they are not "duplicate" errors.
    # …BUT a canonical only legitimately resolves duplication when the canonical
    # TARGET actually carries the same value. Blanket-excluding every page that
    # canonicalises away produced a FALSE PASS on sites that point every page at
    # an unrelated URL: 21 pages all titled "Grant Writing INC." were reported as
    # "All titles are unique" because 21 of 22 pages were dropped from the scope.
    by_canon_url = {canon_url(p.url): p for p in html_pages}

    def dedup_scope(field):
        out = []
        for p in html_pages:
            t = canonical_target(p, start_url)
            if t is not None and canon_url(t) != canon_url(p.url):
                tgt = by_canon_url.get(canon_url(t))
                if tgt is not None and getattr(tgt, field) == getattr(p, field):
                    continue      # true duplicate — the canonical handles it
            out.append(p)
        return out


    total_links = max(1, sum(p.link_count for p in onpage))
    total_imgs = sum(p.images_total for p in onpage)

    # --- ERROR: broken links -----------------------------------------------
    # One report row PER (page → broken target) pair, so the report lists every
    # page that contains a broken link (this is what "22 broken links" means and
    # how Semrush reports it: Page URL | Broken Link | Status).
    # Only links CONFIRMED broken (HTTP 404/410 or a hard DNS/connection
    # failure) are reported as broken. Links that could NOT be verified —
    # bot-blocked or transient responses (403/429/5xx/timeout/SSL) — are kept
    # separately and reported in their own "could not verify" section with the
    # reason, never asserted as broken.
    broken_items, broken_instances, soft404 = [], 0, 0
    broken_urls, broken_pages = set(), set()
    unverified_map = {}          # target -> (reason, sorted sources, internal?)
    for target, sources in link_sources.items():
        pg = pages.get(target)
        internal = same_site(start_url, target)
        loc = "internal" if internal else "external"
        soft = (pg is not None and pg.ok and is_error_page(pg)
                and canon_url(pg.url) != target)
        if soft:
            note = "soft 404 — link redirects to a not-found page"
        else:
            if target in link_status:
                status, err = link_status[target][0], link_status[target][1]
            elif pg is not None:
                status, err = pg.status, pg.error
            else:
                continue
            kind, vnote = link_verdict(status, err)
            if kind == "ok":
                continue
            if kind == "unverified":
                unverified_map[target] = (verify_reason(status, err),
                                          sorted(sources), internal)
                continue
            note = vnote
        soft404 += 1 if soft else 0
        broken_urls.add(target)
        for src in sorted(sources):
            broken_instances += 1
            broken_pages.add(src)
            broken_items.append({
                "url": src,     # the PAGE that contains the broken link
                "note": f"broken link → {truncate(target, 75)}  ({note}, {loc})"})
    broken_items.sort(key=lambda b: (b["note"], b["url"]))
    if broken_instances:
        summary = (f"{broken_instances} broken link(s) on {len(broken_pages)} page(s) "
                   f"({len(broken_urls)} unique broken URL(s)"
                   + (f", incl. {soft404} soft 404" if soft404 else "") + ")")
    else:
        summary = "No broken links detected"
    if not external_checked:
        summary += " (external links skipped)"
    findings.append(finding("broken_links", "Broken internal & external links",
                            "error", summary, broken_items,
                            failed=broken_instances, total=total_links, unit="links"))

    # --- INFO: links that could not be verified (NOT scored) ---------------
    # Reported transparently in their own section with the reason, so the report
    # neither hides them nor falsely claims they are broken.
    if unverified_map:
        uv_items = []
        for target, (reason, srcs, intern) in sorted(unverified_map.items()):
            uv_items.append({
                "url": target,
                "note": f"{reason}  ({'internal' if intern else 'external'}; "
                        f"linked from {len(srcs)} page(s))",
                "sources": srcs[:5]})
        uv_items.sort(key=lambda x: x["url"])
        findings.append(finding(
            "unverified_links", "Links that could not be verified", "info",
            f"{len(unverified_map)} link(s) could not be confirmed working or "
            f"broken (bot-blocked or transient). NOT counted as broken and NOT "
            f"included in the score — listed for manual review.",
            uv_items, failed=len(unverified_map), total=0))

    # --- ERROR: 4xx/5xx status pages ---------------------------------------
    bad_status = [{"url": p.url, "note": f"HTTP {p.status}"}
                  for p in pages.values()
                  if p.status is not None and p.status >= 400]
    findings.append(finding("http_status", "Pages returning 4xx/5xx status codes",
                            "error",
                            f"{len(bad_status)} page(s) returned a 4xx/5xx status"
                            if bad_status else "No 4xx/5xx pages",
                            bad_status, failed=len(bad_status),
                            total=len(pages), unit="pages"))

    # --- ERROR: duplicated titles ------------------------------------------
    title_scope = dedup_scope("title")
    titles = defaultdict(list)
    for p in title_scope:
        if p.title:
            titles[p.title].append(p.url)
    dup_titles = {t: u for t, u in titles.items() if len(u) > 1}
    dup_title_pages = sum(len(u) for u in dup_titles.values())
    items = [{"url": urls[0], "note": f'"{truncate(t)}" — on {len(urls)} pages',
              "sources": urls[:25]} for t, urls in dup_titles.items()]
    findings.append(finding("dup_titles", "Duplicate title tags", "error",
                            f"{dup_title_pages} page(s) share a duplicate title "
                            f"({len(dup_titles)} duplicated title(s))"
                            if dup_titles else "All titles are unique",
                            items, failed=dup_title_pages,
                            total=max(1, len(title_scope))))

    # --- ERROR: duplicated meta descriptions -------------------------------
    desc_scope = dedup_scope("meta_desc")
    descs = defaultdict(list)
    for p in desc_scope:
        if p.meta_desc:
            descs[p.meta_desc].append(p.url)
    dup_desc = {d: u for d, u in descs.items() if len(u) > 1}
    dup_desc_pages = sum(len(u) for u in dup_desc.values())
    items = [{"url": urls[0], "note": f'"{truncate(d)}" — on {len(urls)} pages',
              "sources": urls[:25]} for d, urls in dup_desc.items()]
    findings.append(finding("dup_meta", "Duplicate meta descriptions", "error",
                            f"{dup_desc_pages} page(s) share a duplicate meta "
                            f"description ({len(dup_desc)} duplicated)"
                            if dup_desc else "All meta descriptions are unique",
                            items, failed=dup_desc_pages,
                            total=max(1, len(desc_scope))))

    # --- ERROR: duplicate content ------------------------------------------
    # Cluster ALL indexable pages, then drop only those pages whose rel=canonical
    # points at a URL inside the SAME cluster — that is the case a canonical really
    # resolves. A canonical aimed at an unrelated page leaves the duplication real.
    raw_clusters = find_duplicate_content(html_pages)
    dup_clusters = []
    for cl in raw_clusters:
        keys = {canon_url(u) for u in cl}
        kept = []
        for u in cl:
            p = by_canon_url.get(canon_url(u))
            t = canonical_target(p, start_url) if p is not None else None
            if (t is not None and canon_url(t) != canon_url(u)
                    and canon_url(t) in keys):
                continue
            kept.append(u)
        if len(kept) > 1:
            dup_clusters.append(kept)
    dup_content_pages = sum(len(c) for c in dup_clusters)
    items = [{"url": c[0], "note": f"near-duplicate body content on {len(c)} pages",
              "sources": c[:25]} for c in dup_clusters]
    findings.append(finding("dup_content", "Duplicate content", "error",
                            f"{dup_content_pages} page(s) in {len(dup_clusters)} "
                            f"near-duplicate cluster(s)"
                            if dup_clusters else "No duplicate content detected",
                            items, failed=dup_content_pages,
                            total=max(1, len(html_pages))))

    # --- ERROR: incorrect pages in sitemap ---------------------------------
    # Judge each sitemap URL EXACTLY as listed. Two bugs made this wildly
    # under-report (10 found against a true 120 on americanwebbuilders.com):
    #   1. Sitemap URLs that were never crawled AND never linked were skipped
    #      entirely — no status was ever fetched for them. main() now adds every
    #      same-site sitemap URL to the link-verification set.
    #   2. `pages.get(canon_url(u))` folds trailing-slash variants together, so a
    #      sitemap entry "/x" resolved to the crawled "/x/" page (HTTP 200, no
    #      redirect chain) and looked healthy — while "/x" itself 301s. The exact
    #      URL's own status now wins over the canon-folded page lookup.
    sm_bad = []
    sm_same = [u for u in sitemap["urls"] if same_site(start_url, u)]
    for u in sm_same:
        page_obj = pages.get(canon_url(u))
        exact = link_status.get(u)
        if exact is None and page_obj is not None and page_obj.requested_url == u:
            exact = "page"
        if exact is not None and exact != "page":
            st, er, fh, red = exact
            if red:
                sm_bad.append({"url": u, "note": "sitemap URL redirects elsewhere"})
                continue
            if st in BROKEN_STATUSES:
                sm_bad.append({"url": u, "note": f"sitemap URL is broken (HTTP {st})"})
                continue
        elif page_obj is not None:
            if page_obj.redirect_chain:
                sm_bad.append({"url": u, "note": "sitemap URL redirects elsewhere"})
                continue
            if page_obj.status in BROKEN_STATUSES:
                sm_bad.append({"url": u, "note": f"sitemap URL is broken (HTTP {page_obj.status})"})
                continue
        # URL itself resolves 200 — is the page it serves self-canonical?
        if page_obj is not None and canonicalizes_away(page_obj, start_url):
            sm_bad.append({"url": u, "note": "sitemap URL canonicalises to a different "
                                             "page — non-canonical URLs should not be in the sitemap"})
    findings.append(finding("sitemap_incorrect", "Incorrect pages in sitemap.xml",
                            "error",
                            f"{len(sm_bad)} sitemap URL(s) redirect, are broken, or are non-canonical"
                            if sm_bad else "Sitemap URLs are clean",
                            sm_bad, failed=len(sm_bad),
                            total=max(1, len(sm_same)), unit="sitemap URLs"))

    # --- WARNING: unminified JS/CSS ----------------------------------------
    # Counted PER PAGE-REFERENCE (one unminified file loaded on N pages = N
    # issues), matching how Semrush reports these, with the unique-file count
    # shown alongside.
    unmin_files = {u for u, flag in resource_status.items() if flag}
    unmin_refs = total_refs = 0
    for p in onpage:
        for a in p.assets:
            if re.search(r"\.(css|js)(\?|#|$)", a, re.I):
                total_refs += 1
                if a in unmin_files:
                    unmin_refs += 1
    findings.append(finding("unminified", "Unminified JavaScript and CSS files",
                            "warning",
                            f"{unmin_refs} unminified reference(s) across "
                            f"{len(unmin_files)} unique file(s)"
                            if unmin_refs else
                            (f"All {len(resource_status)} checked JS/CSS files are minified"
                             if resource_status else "No JS/CSS files checked"),
                            [{"url": u, "note": "not minified"} for u in sorted(unmin_files)[:50]],
                            failed=unmin_refs, total=max(1, total_refs), unit="references"))

    # --- WARNING: images without alt ---------------------------------------
    alt_items, missing_alt = [], 0
    for p in onpage:
        missing_alt += p.images_missing_alt
        if p.images_missing_alt:
            alt_items.append({"url": p.url,
                              "note": f"{p.images_missing_alt}/{p.images_total} images missing alt",
                              "sources": p.missing_alt_samples})
    findings.append(finding("image_alt", "Images without alt attributes",
                            "warning",
                            f"{missing_alt} image(s) missing alt across {len(alt_items)} page(s)"
                            if missing_alt else
                            (f"All {total_imgs} images have alt text" if total_imgs
                             else "No images found"),
                            alt_items, failed=missing_alt,
                            total=max(1, total_imgs), unit="images"))

    # --- WARNING: low text-HTML ratio --------------------------------------
    low_ratio = []
    for p in onpage:
        if p.html_len:
            ratio = p.text_len / p.html_len
            if ratio < TEXT_HTML_MIN:
                low_ratio.append({"url": p.url,
                                  "note": f"text/HTML ratio {ratio*100:.1f}% "
                                          f"(min {TEXT_HTML_MIN*100:.0f}%)"})
    findings.append(finding("text_ratio", "Low text-to-HTML ratio", "warning",
                            f"{len(low_ratio)} page(s) below {TEXT_HTML_MIN*100:.0f}% text/HTML"
                            if low_ratio else "Text-to-HTML ratio is healthy",
                            low_ratio, failed=len(low_ratio), total=n_pages))

    # --- WARNING: nofollow internal links ----------------------------------
    nofollow_instances = sum(len(p.nofollow_internal) for p in onpage)
    nf_items = [{"url": p.url, "note": f"{len(p.nofollow_internal)} nofollow internal link(s)"}
                for p in onpage if p.nofollow_internal]
    findings.append(finding("nofollow", "Outgoing internal links with nofollow",
                            "warning",
                            f"{nofollow_instances} internal link(s) marked nofollow"
                            if nofollow_instances else "No nofollowed internal links",
                            nf_items, failed=nofollow_instances,
                            total=total_links, unit="links"))

    # --- WARNING: missing H1 -----------------------------------------------
    missing_h1 = [{"url": p.url, "note": "no H1 tag"} for p in onpage if len(p.h1s) == 0]
    findings.append(finding("missing_h1", "Pages without an H1 heading", "warning",
                            f"{len(missing_h1)} page(s) have no H1"
                            if missing_h1 else "Every page has an H1",
                            missing_h1, failed=len(missing_h1), total=n_pages))

    # --- WARNING: missing meta description ---------------------------------
    missing_desc = [p.url for p in onpage if not p.meta_desc]
    md_items = [{"url": u, "note": "missing meta description"} for u in missing_desc]
    findings.append(finding("missing_meta", "Pages without a meta description",
                            "warning",
                            f"{len(missing_desc)} page(s) missing a meta description"
                            if missing_desc else "Every page has a meta description",
                            md_items, failed=len(missing_desc), total=n_pages))

    # --- WARNING: missing title --------------------------------------------
    missing_title = [p.url for p in onpage if not p.title]
    mt_items = [{"url": u, "note": "missing <title>"} for u in missing_title]
    findings.append(finding("missing_title", "Pages without a title tag", "warning",
                            f"{len(missing_title)} page(s) missing a <title>"
                            if missing_title else "Every page has a title",
                            mt_items, failed=len(missing_title), total=n_pages))

    # --- WARNING: low word count -------------------------------------------
    low_wc = [{"url": p.url, "note": f"{p.word_count} words (min {LOW_WORDS})"}
              for p in onpage if p.word_count < LOW_WORDS]
    findings.append(finding("low_word_count", "Pages with low word count", "warning",
                            f"{len(low_wc)} page(s) under {LOW_WORDS} words"
                            if low_wc else "No thin pages by word count",
                            low_wc, failed=len(low_wc), total=n_pages))

    # --- WARNING: title length ---------------------------------------------
    title_len = []
    for p in onpage:
        if not p.title:
            continue
        n = len(p.title)
        if n > TITLE_MAX:
            title_len.append({"url": p.url, "note": f"title is {n} chars (max {TITLE_MAX})"})
        elif n < TITLE_MIN:
            title_len.append({"url": p.url, "note": f"title is only {n} chars (min {TITLE_MIN})"})
    findings.append(finding("title_length", "Title tags too long or too short",
                            "warning",
                            f"{len(title_len)} title(s) outside {TITLE_MIN}-{TITLE_MAX} chars"
                            if title_len else "Title lengths are healthy",
                            title_len, failed=len(title_len), total=n_pages))

    # --- WARNING: meta description length ----------------------------------
    desc_len = []
    for p in onpage:
        if not p.meta_desc:
            continue
        n = len(p.meta_desc)
        if n > DESC_MAX:
            desc_len.append({"url": p.url, "note": f"meta description is {n} chars (max {DESC_MAX})"})
        elif n < DESC_MIN:
            desc_len.append({"url": p.url, "note": f"meta description is only {n} chars (min {DESC_MIN})"})
    findings.append(finding("desc_length", "Meta descriptions too long or too short",
                            "warning",
                            f"{len(desc_len)} meta description(s) outside {DESC_MIN}-{DESC_MAX} chars"
                            if desc_len else "Meta description lengths are healthy",
                            desc_len, failed=len(desc_len), total=n_pages))

    # --- WARNING: HTTPS page links to HTTP ---------------------------------
    https_http = sum(p.http_from_https for p in onpage)
    hh_items = [{"url": p.url, "note": f"{p.http_from_https} link(s) to HTTP URLs"}
                for p in onpage if p.http_from_https]
    findings.append(finding("https_to_http", "Links from HTTPS pages to HTTP URLs",
                            "warning",
                            f"{https_http} HTTPS→HTTP link(s)"
                            if https_http else "No HTTPS→HTTP links",
                            hh_items, failed=https_http, total=total_links, unit="links"))

    # --- WARNING: mixed content --------------------------------------------
    mixed = sum(p.mixed_content for p in onpage)
    mc_items = [{"url": p.url, "note": f"{p.mixed_content} resource(s) loaded over HTTP"}
                for p in onpage if p.mixed_content]
    findings.append(finding("mixed_content", "Mixed content (HTTP resources on HTTPS)",
                            "warning",
                            f"{mixed} HTTP resource(s) loaded on HTTPS pages"
                            if mixed else "No mixed content",
                            mc_items, failed=mixed, total=total_links, unit="resources"))

    # --- WARNING: missing viewport (mobile) --------------------------------
    no_viewport = [{"url": p.url, "note": "no <meta name=viewport>"}
                   for p in onpage if not p.has_viewport]
    findings.append(finding("viewport", "Pages without a viewport meta tag", "warning",
                            f"{len(no_viewport)} page(s) missing a viewport meta tag"
                            if no_viewport else "Every page declares a viewport",
                            no_viewport, failed=len(no_viewport), total=n_pages))

    # --- NOTICE: resources formatted as page link --------------------------
    res_instances = sum(p.resource_link_count for p in onpage)
    res_items = [{"url": p.url, "note": f"{p.resource_link_count} resource link(s)"}
                 for p in onpage if p.resource_link_count]
    findings.append(finding("resource_links", "Resources formatted as page link",
                            "notice",
                            f"{res_instances} link(s) point at a resource file"
                            if res_instances else "No resource-as-page links",
                            res_items, failed=res_instances, total=total_links, unit="links"))

    # --- NOTICE: links with no anchor text ---------------------------------
    # Counted per instance (like every other check), but unnamed links are almost
    # always a handful of TEMPLATE links — a logo <img> with no alt, social icons
    # — repeated on every page. Report the unique-target count alongside the
    # instance count so "970" reads as "5 template links to fix", not 970 jobs.
    empty_anchor = sum(p.empty_anchor for p in onpage)
    ea_targets = Counter(u for p in onpage for u in p.empty_anchor_urls)
    ea_items = [{"url": u, "note": f"linked with no anchor text on {c} page(s)"}
                for u, c in ea_targets.most_common()]
    findings.append(finding("empty_anchor", "Links with no anchor text", "notice",
                            f"{empty_anchor} link(s) have no anchor text "
                            f"across {len(ea_targets)} unique link target(s)"
                            if empty_anchor else "All links have anchor text",
                            ea_items, failed=empty_anchor, total=total_links, unit="links"))

    # --- NOTICE: non-descriptive anchor text -------------------------------
    nondesc = sum(p.nondesc_anchor for p in onpage)
    nd_items = [{"url": p.url, "note": f"{p.nondesc_anchor} non-descriptive anchor(s)"}
                for p in onpage if p.nondesc_anchor]
    findings.append(finding("nondesc_anchor", "Non-descriptive anchor text", "notice",
                            f"{nondesc} link(s) use generic anchor text"
                            if nondesc else "Anchor text is descriptive",
                            nd_items, failed=nondesc, total=total_links, unit="links"))

    # --- NOTICE: permanent (301) redirects ---------------------------------
    # INTERNAL links pointing at a 301 only — a link on the site sending users
    # (and link equity) through a redirect. External sites redirecting their own
    # URLs (social profiles, etc.) is not this site's SEO concern, so it is
    # excluded, matching Semrush's "URLs with a permanent redirect" report.
    perm_items, perm = [], 0
    #
    # Counted PER (page → redirecting target) pair, like broken links, because
    # the cost is paid on every page carrying the link. Two earlier bugs, both
    # found by an independent re-crawl of americanwebbuilders.com:
    #   1. It counted each redirecting URL ONCE, so the number was neither the
    #      unique-URL count nor the link count — 151, against a true 9 unique
    #      URLs / 548 link instances.
    #   2. 143 of those 151 were .webp/.css/.js files linked via <a>, not pages.
    #      Resource files are already reported by "resources formatted as page
    #      link"; a 301 on an image is not what this check is about.
    # Uses raw_link_sources (links AS WRITTEN) rather than the canon_url-folded
    # link_sources, so www/non-www and trailing-slash variants stay distinct —
    # that folding hid 154 links to the redirecting non-www homepage.
    redir_targets = {}
    for target, sources in (raw_link_sources or link_sources).items():
        if not same_site(start_url, target) or RESOURCE_EXT.search(target):
            continue
        is301, dest = False, ""
        if target in link_status:
            st, er, fh, red = link_status[target]
            if red and fh == 301:
                is301 = True
        else:
            pg = pages.get(canon_url(target))
            if (pg is not None and pg.requested_url == target
                    and any(st == 301 for st, _ in pg.redirect_chain)):
                is301, dest = True, pg.url
        if is301:
            redir_targets[target] = (len(sources), dest)
    perm = sum(n for n, _ in redir_targets.values())
    perm_items = [{"url": u, "note": (f"301 → {truncate(dest)} — linked from "
                                      f"{n} page(s)" if dest else
                                      f"301 permanent redirect — linked from {n} page(s)")}
                  for u, (n, dest) in sorted(redir_targets.items(),
                                             key=lambda kv: -kv[1][0])]
    findings.append(finding("permanent_redirects", "Permanent (301) redirects",
                            "notice",
                            f"{perm} internal link(s) point at a 301 redirect "
                            f"across {len(redir_targets)} unique redirecting URL(s)"
                            if perm else "No internal permanent redirects",
                            perm_items, failed=perm, total=total_links, unit="links"))

    # --- NOTICE: multiple H1 -----------------------------------------------
    multi_h1 = [{"url": p.url, "note": f"{len(p.h1s)} H1 tags"}
                for p in onpage if len(p.h1s) > 1]
    findings.append(finding("multiple_h1", "Pages with more than one H1", "notice",
                            f"{len(multi_h1)} page(s) have multiple H1s"
                            if multi_h1 else "No pages with multiple H1s",
                            multi_h1, failed=len(multi_h1), total=n_pages))

    # --- NOTICE: multiple title tags ---------------------------------------
    multi_title = [{"url": p.url, "note": f"{p.title_count} <title> tags"}
                   for p in onpage if p.title_count > 1]
    findings.append(finding("multiple_title", "Pages with more than one title tag",
                            "notice",
                            f"{len(multi_title)} page(s) have multiple <title> tags"
                            if multi_title else "No pages with multiple titles",
                            multi_title, failed=len(multi_title), total=n_pages))

    # --- NOTICE: missing charset declaration -------------------------------
    no_charset = [{"url": p.url, "note": "no charset declaration"}
                  for p in onpage if not p.has_charset]
    findings.append(finding("charset", "Pages without a charset declaration", "notice",
                            f"{len(no_charset)} page(s) missing a charset declaration"
                            if no_charset else "Every page declares a charset",
                            no_charset, failed=len(no_charset), total=n_pages))

    # --- NOTICE: missing doctype -------------------------------------------
    no_doctype = [{"url": p.url, "note": "no <!doctype> declaration"}
                  for p in onpage if not p.has_doctype]
    findings.append(finding("doctype", "Pages without a doctype", "notice",
                            f"{len(no_doctype)} page(s) missing a doctype"
                            if no_doctype else "Every page declares a doctype",
                            no_doctype, failed=len(no_doctype), total=n_pages))

    # --- NOTICE: pages with only one incoming internal link ----------------
    if crawl_complete:
        weak = [p.url for p in onpage
                if canon_url(p.url) != seed
                and len(link_sources.get(canon_url(p.url), set())) == 1]
        findings.append(finding("weak_linking",
                                "Pages with only one incoming internal link",
                                "notice",
                                f"{len(weak)} page(s) have a single inbound internal link"
                                if weak else "No weakly-linked pages",
                                [{"url": u, "note": "only 1 incoming internal link"} for u in weak],
                                failed=len(weak), total=n_pages))
    else:
        findings.append(finding("weak_linking",
                                "Pages with only one incoming internal link",
                                "notice",
                                f"Skipped — crawl hit the {max_pages}-page cap; "
                                "internal link graph is incomplete",
                                [], failed=0, total=0))

    # --- NOTICE: orphan pages ----------------------------------------------
    if not crawl_complete:
        findings.append(finding("orphans", "Orphan pages", "notice",
                                f"Skipped — crawl hit the {max_pages}-page cap. "
                                "Re-run with --max-pages >= total pages to detect orphans.",
                                [], failed=0, total=0))
    elif sitemap["count"]:
        linked = set(link_sources.keys()) | set(pages.keys()) | {seed}
        sm_canon = {canon_url(u): u for u in sitemap["urls"] if same_site(start_url, u)}
        orphans = sorted(orig for c, orig in sm_canon.items() if c not in linked)
        findings.append(finding("orphans", "Orphan pages", "notice",
                                f"{len(orphans)} sitemap URL(s) not internally linked"
                                if orphans else "No orphan pages",
                                [{"url": u, "note": "in sitemap but no internal link points to it"}
                                 for u in orphans],
                                failed=len(orphans), total=max(1, len(sm_canon)),
                                unit="sitemap URLs"))
    else:
        findings.append(finding("orphans", "Orphan pages", "notice",
                                "No XML sitemap to compare against", [],
                                failed=0, total=0))

    # --- NOTICE: non-indexable pages ---------------------------------------
    nonindex = []
    rp = robots.get("parser")
    for p in html_pages:
        reasons = []
        if "noindex" in (p.robots_meta or ""):
            reasons.append("meta robots noindex")
        if "noindex" in (p.x_robots or "").lower():
            reasons.append("X-Robots-Tag noindex")
        if rp is not None:
            try:
                if not rp.can_fetch(UA, p.url):
                    reasons.append("blocked by robots.txt")
            except Exception:
                pass
        if reasons:
            nonindex.append({"url": p.url, "note": "; ".join(reasons)})
    findings.append(finding("non_indexable", "Non-indexable pages", "notice",
                            f"{len(nonindex)} non-indexable page(s) — confirm intentional"
                            if nonindex else "All crawled pages are indexable",
                            nonindex, failed=len(nonindex), total=len(html_pages) or 1))

    # --- NOTICE: canonical tag issues --------------------------------------
    canon_issues, canon_away_items = [], []
    for p in html_pages:
        if not p.canonicals:
            canon_issues.append({"url": p.url, "note": "no canonical tag (recommended)"})
            continue
        if len(p.canonicals) > 1:
            canon_issues.append({"url": p.url,
                                 "note": f"conflicting: {len(p.canonicals)} canonical tags"})
        href = p.canonicals[0]
        if not re.match(r"^https?://", href, re.I):
            target = urllib.parse.urljoin(p.url, href)
            canon_issues.append({"url": p.url,
                                 "note": f"relative canonical '{truncate(href)}' — use absolute URL"})
        else:
            target = href
        if not same_site(start_url, target):
            canon_issues.append({"url": p.url,
                                 "note": f"canonical points off-site: {truncate(target)}"})
        elif canon_url(target) != canon_url(p.url):
            canon_away_items.append({"url": p.url,
                                     "note": f"canonical → {truncate(target)}"})
    # This row covers only MALFORMED canonicals (missing / conflicting / relative /
    # off-site). Cross-page canonicals get their own row below. Never claim the
    # canonicals "look correct" while that row is non-empty — the two rows together
    # were reported as a contradiction ("Canonical tags look correct" next to 20
    # pages canonicalising to the homepage).
    if canon_issues:
        canon_summary = f"{len(canon_issues)} canonical note(s)"
    elif canon_away_items:
        canon_summary = (f"tags are well-formed, but {len(canon_away_items)} page(s) "
                         f"canonicalise to a different URL — see \"Pages "
                         f"canonicalised to a different URL\"")
    else:
        canon_summary = "Canonical tags look correct"
    findings.append(finding("canonical", "Canonical tag issues", "notice",
                            canon_summary,
                            canon_issues, failed=len(canon_issues),
                            total=len(html_pages) or 1))

    # --- WARNING: pages canonicalised to a different URL -------------------
    # These pages tell search engines "index that other URL instead of me", so
    # they will not rank on their own. A handful (param variants) is normal; a
    # large share — e.g. every page pointing at the homepage — is a critical,
    # site-wide indexation problem that Semrush reports here.
    findings.append(finding("canonicalized",
                            "Pages canonicalised to a different URL", "warning",
                            f"{len(canon_away_items)} page(s) point rel=canonical at "
                            f"another URL and won't be indexed on their own"
                            if canon_away_items else "No pages canonicalise away",
                            canon_away_items, failed=len(canon_away_items),
                            total=len(html_pages) or 1))

    # --- NOTICE: slow-loading pages ----------------------------------------
    timings = [p.elapsed for p in html_pages if p.elapsed]
    slow = [{"url": p.url, "note": f"{p.elapsed}s (threshold {slow_threshold}s)"}
            for p in sorted(html_pages, key=lambda x: -x.elapsed)
            if p.elapsed > slow_threshold]
    avg = round(sum(timings) / len(timings), 2) if timings else 0
    findings.append(finding("slow_pages", "Slow-loading pages", "notice",
                            f"{len(slow)} slow page(s); avg load {avg}s"
                            if slow else f"No slow pages (avg load {avg}s)",
                            slow, failed=len(slow), total=len(html_pages) or 1))

    # --- NOTICE: HSTS -------------------------------------------------------
    seed_page = pages.get(seed) or next(iter(pages.values()), None)
    hsts_ok = bool(seed_page and seed_page.hsts)
    findings.append(finding("hsts", "HSTS support", "notice",
                            "HSTS (Strict-Transport-Security) is enabled" if hsts_ok
                            else "HSTS (Strict-Transport-Security) is not enabled",
                            [] if hsts_ok else
                            [{"url": start_url, "note": "no Strict-Transport-Security header"}],
                            failed=0 if hsts_ok else 1, total=1))

    # --- WARNING: www / non-www duplicate host -----------------------------
    if host_variants is not None:
        dup = host_variants.get("duplicate")
        findings.append(finding("host_duplicate",
                                "www and non-www both serve content (duplicate host)",
                                "warning",
                                "Both the www and non-www hosts return HTTP 200 without "
                                "redirecting to one canonical host — every URL exists twice. "
                                "Add a 301 redirect to your preferred host."
                                if dup else "A single canonical host is enforced",
                                [{"url": host_variants["nonwww"], "note": "serves 200"},
                                 {"url": host_variants["www"], "note": "serves 200"}]
                                if dup else [],
                                failed=1 if dup else 0, total=1))

    # --- NOTICE: robots.txt / sitemap health -------------------------------
    sr_items, sr_fail = [], 0
    if not robots["exists"]:
        sr_items.append({"url": urllib.parse.urljoin(start_url, "/robots.txt"),
                         "note": f"robots.txt missing or empty (HTTP {robots['status']})"})
        sr_fail += 1
    else:
        if robots["blocks_all"]:
            sr_items.append({"url": "robots.txt",
                             "note": "Disallow: / blocks the ENTIRE site for all bots"})
            sr_fail += 1
        if not robots["sitemaps"]:
            sr_items.append({"url": "robots.txt",
                             "note": "no Sitemap: directive in robots.txt (recommended)"})
            sr_fail += 1
    for issue in robots.get("issues", []):
        sr_items.append({"url": "robots.txt", "note": issue})
        sr_fail += 1
    if sitemap["count"] == 0:
        sr_items.append({"url": urllib.parse.urljoin(start_url, "/sitemap.xml"),
                         "note": "no valid XML sitemap found"})
        sr_fail += 1
    for issue in sitemap.get("issues", []):
        sr_items.append({"url": "sitemap", "note": issue})
        sr_fail += 1
    tier = "error" if (robots.get("blocks_all")) else "notice"
    findings.append(finding("sitemap_robots", "robots.txt & sitemap health", tier,
                            f"robots.txt: {'present' if robots['exists'] else 'missing'}; "
                            f"sitemap: {sitemap['count']} URL(s) in "
                            f"{len(sitemap['found'])} file(s)",
                            sr_items, failed=sr_fail, total=max(1, sr_fail + 4)))

    return findings


# ===========================================================================
# Scoring — Semrush-style weighted pass-ratio
# ===========================================================================
def site_health(findings):
    """Site Health = weighted average across all APPLICABLE checks of the
    fraction of units that PASS, weighted by severity tier. Mirrors Semrush:
    proportional (frequency of the issue), severity-weighted (errors > warnings
    > notices), and driven by checks-run rather than a flat per-category
    penalty. Passing checks hold the score up; each failing check drags its own
    term down in proportion to how many units are affected."""
    num = den = 0.0
    for f in findings:
        total = f.get("total", 0)
        w = TIER_WEIGHT.get(f["tier"])       # "info" tier -> None -> never scored
        if w is None or total <= 0:          # not applicable / informational
            continue
        fail_frac = min(1.0, f["failed"] / total)
        pass_frac = 1.0 - (fail_frac ** PENALTY_EXP)
        num += w * pass_frac
        den += w
    return round(100 * num / den) if den else 100


# ===========================================================================
# Output — console
# ===========================================================================
def _group(findings):
    groups = {"error": [], "warning": [], "notice": [], "info": [], "passed": []}
    for f in findings:
        groups.setdefault(f["display"], []).append(f)
    return groups


def print_console(start_url, pages, findings, score, stats):
    ok_pages = sum(1 for p in pages.values() if p.ok)
    groups = _group(findings)
    print()
    print("=" * 70)
    print("  TECHNICAL SEO AUDIT  (Semrush-style)")
    print("=" * 70)
    print(f"  Site         : {start_url}")
    print(f"  Pages crawled: {len(pages)}  ({ok_pages} returned HTTP 200)"
          + ("  [JS-rendered]" if stats.get("rendered") else ""))
    print(f"  Links checked: {stats['links_checked']}"
          + ("  [external skipped]" if not stats["external_checked"] else ""))
    print(f"  JS/CSS files : {stats.get('resources_checked', 0)} checked for minification")
    print(f"  SITE HEALTH  : {score}/100")
    print(f"  Issues       : {len(groups['error'])} errors, "
          f"{len(groups['warning'])} warnings, {len(groups['notice'])} notices, "
          f"{len(groups['passed'])} passed")
    if groups["info"]:
        n_uv = sum(len(f["items"]) for f in groups["info"])
        print(f"  Unverified   : {n_uv} link(s) could not be checked "
              f"(listed separately, not scored)")
    print("=" * 70)

    if stats.get("content_warning"):
        print()
        print("  " + "!" * 66)
        print("  CRITICAL: this page has almost no indexable content.")
        print(f"  Cause: {stats['content_warning']}")
        print("  A search engine sees an essentially empty page; the score is capped.")
        print("  " + "!" * 66)

    for tier in ("error", "warning", "notice", "info", "passed"):
        fs = groups[tier]
        if not fs:
            continue
        print(f"\n  {'-'*3} {SECTION_TITLE[tier]} ({len(fs)}) {'-'*3}\n")
        for f in fs:
            tag = TIER_LABEL[f["display"]]
            print(f"  [{tag:7}] {f['name']}")
            print(f"            {f['summary']}")
            for item in f["items"][:6]:
                print(f"              - {truncate(item['url'], 76)}")
                if item.get("note"):
                    print(f"                  {item['note']}")
                for s in item.get("sources", [])[:2]:
                    print(f"                  on: {truncate(s, 70)}")
            if len(f["items"]) > 6:
                print(f"              ... and {len(f['items']) - 6} more")
            print()


# ===========================================================================
# Output — HTML
# ===========================================================================
def build_html(start_url, pages, findings, score, stats):
    def esc(s):
        return html_lib.escape(str(s))

    color = "#16a34a" if score >= 80 else "#d97706" if score >= 55 else "#dc2626"
    ok_pages = sum(1 for p in pages.values() if p.ok)
    groups = _group(findings)

    def card(f):
        disp = f["display"]
        rows = []
        for item in f["items"][:200]:
            src = ""
            if item.get("sources"):
                src = ("<div class='src'>" +
                       "<br>".join("&#8627; " + esc(truncate(s, 90))
                                   for s in item["sources"]) + "</div>")
            rows.append(
                f"<div class='item'><a href='{esc(item['url'])}' target='_blank'>"
                f"{esc(truncate(item['url'], 90))}</a>"
                f"<div class='note'>{esc(item.get('note', ''))}</div>{src}</div>")
        if len(f["items"]) > 200:
            rows.append(f"<div class='note'>… and {len(f['items'])-200} more</div>")
        body = "".join(rows) if rows else "<div class='ok'>&#10003; No issues</div>"
        return f"""
        <div class="card">
          <div class="cardhead">
            <span class="badge" style="background:{TIER_COLOR[disp]}">{TIER_LABEL[disp]}</span>
            <h2>{esc(f['name'])}</h2>
          </div>
          <div class="summary">{esc(f['summary'])}</div>
          <div class="items">{body}</div>
        </div>"""

    sections = []
    for tier in ("error", "warning", "notice", "info", "passed"):
        fs = groups[tier]
        if not fs:
            continue
        sections.append(f"<h3 class='sec' style='color:{TIER_COLOR[tier]}'>"
                        f"{SECTION_TITLE[tier]} ({len(fs)})</h3>")
        sections.append("".join(card(f) for f in fs))

    crit_banner = ""
    if stats.get("content_warning"):
        crit_banner = (
            '<div class="crit"><b>&#9888; CRITICAL — no indexable content.</b> '
            f'{esc(stats["content_warning"])}. The score is capped.</div>')

    return f"""<!doctype html><html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Technical SEO Audit — {esc(start_url)}</title>
<style>
  body{{font:15px/1.6 -apple-system,Segoe UI,Roboto,Arial,sans-serif;color:#0f172a;
       max-width:940px;margin:36px auto;padding:0 20px;background:#f8fafc}}
  h1{{font-size:23px;margin:0 0 4px}} .url{{color:#475569;word-break:break-all}}
  .hero{{display:flex;align-items:center;gap:26px;background:#fff;border:1px solid #e2e8f0;
        border-radius:14px;padding:22px;margin:18px 0;box-shadow:0 1px 3px rgba(0,0,0,.05)}}
  .score{{font-size:52px;font-weight:800;color:{color};line-height:1}}
  .meta{{color:#64748b;font-size:13px}} .stat{{display:inline-block;margin-right:18px}}
  .sec{{margin:26px 0 6px;font-size:15px;letter-spacing:.5px}}
  .card{{background:#fff;border:1px solid #e2e8f0;border-radius:14px;padding:16px 20px;
        margin:12px 0;box-shadow:0 1px 3px rgba(0,0,0,.05)}}
  .cardhead{{display:flex;align-items:center;gap:10px}}
  .card h2{{font-size:16px;margin:0}}
  .badge{{display:inline-block;color:#fff;font-size:11px;font-weight:700;padding:3px 9px;
         border-radius:6px}}
  .summary{{color:#334155;font-size:14px;margin:8px 0 4px}}
  .items{{margin-top:6px;max-height:340px;overflow:auto}}
  .item{{padding:7px 0;border-top:1px solid #f1f5f9}}
  .item a{{color:#2563eb;text-decoration:none;font-size:13px;word-break:break-all}}
  .note{{color:#64748b;font-size:12.5px;margin-top:2px}}
  .src{{color:#94a3b8;font-size:12px;margin-top:2px}}
  .ok{{color:#16a34a;font-size:13px}}
  .crit{{background:#fef2f2;border:1px solid #fecaca;border-left:5px solid #dc2626;
        border-radius:10px;padding:14px 18px;margin:16px 0}}
  .crit b{{color:#dc2626}}
</style></head><body>
  <h1>Technical SEO Audit Report</h1>
  <div class="url">{esc(start_url)}</div>
  {crit_banner}
  <div class="hero">
    <div><div class="score">{score}</div><div class="meta">/ 100 &nbsp;Site Health</div></div>
    <div class="meta">
      <span class="stat">&#128196; {len(pages)} pages crawled{' &#183; JS-rendered' if stats.get('rendered') else ''}</span>
      <span class="stat">&#9989; {ok_pages} OK (200)</span>
      <span class="stat">&#128279; {stats['links_checked']} links checked</span>
      <span class="stat" style="color:#dc2626">&#10007; {len(groups['error'])} errors</span>
      <span class="stat" style="color:#d97706">&#9888; {len(groups['warning'])} warnings</span>
      <span class="stat" style="color:#3b82f6">&#8505; {len(groups['notice'])} notices</span>
    </div>
  </div>
  {''.join(sections)}
  <p class="meta">Generated by seoaudit.py — Semrush-style Technical SEO Audit.
     Site Health = severity-weighted pass ratio across all applicable checks.</p>
</body></html>"""


# ===========================================================================
# Output — Document (.docx / .rtf) : the scheduled audit-report deliverable
# ===========================================================================
# One recommended fixing action per check, keyed by finding id (Semrush-style
# "How to fix"). Falls back to a generic action for anything not listed.
RECOMMENDATIONS = {
    "broken_links": "Update or remove the broken link. Point it to a live URL, "
                    "restore the missing page, or delete the reference.",
    "http_status": "Fix the server error or restore the page. Return 200 for valid "
                   "pages; 301-redirect removed pages to the best alternative.",
    "dup_titles": "Write a unique, descriptive <title> for each page so no two "
                  "pages share the same title.",
    "dup_meta": "Write a unique meta description for each page; do not reuse the "
                "same description across pages.",
    "dup_content": "Consolidate or rewrite near-duplicate pages, or set a "
                   "rel=canonical to the primary version.",
    "sitemap_incorrect": "Remove redirecting, broken, or non-canonical URLs from "
                         "sitemap.xml; list only indexable 200-status canonical URLs.",
    "unminified": "Minify your CSS and JavaScript files (remove whitespace/comments) "
                  "and serve the .min versions to cut page weight.",
    "image_alt": "Add a descriptive alt attribute to every meaningful <img>; use "
                 "alt=\"\" only for purely decorative images.",
    "text_ratio": "Increase real page text and/or reduce heavy inline markup so the "
                  "text-to-HTML ratio rises above ~10%.",
    "nofollow": "Remove rel=\"nofollow\" from internal links — you should pass link "
                "equity between your own pages.",
    "missing_h1": "Add exactly one clear, keyword-relevant <h1> heading to each page.",
    "missing_meta": "Add a unique 70–160 character meta description summarising each "
                    "page for search results.",
    "missing_title": "Add a unique, descriptive <title> tag to every page.",
    "low_word_count": "Expand thin pages with useful, original content (aim well "
                      "above 200 words).",
    "title_length": "Keep titles roughly 10–60 characters so they are neither cut off "
                    "nor too sparse in search results.",
    "desc_length": "Keep meta descriptions ~70–160 characters so they display fully "
                   "in the SERP snippet.",
    "https_to_http": "Change HTTP links to HTTPS so users never leave the secure "
                     "version of the site.",
    "mixed_content": "Serve every resource (images, scripts, styles) over HTTPS to "
                     "remove mixed-content warnings.",
    "viewport": "Add <meta name=\"viewport\" content=\"width=device-width, "
                "initial-scale=1\"> for correct mobile rendering.",
    "canonicalized": "Point each page's rel=canonical at itself (unless it is a true "
                     "duplicate). Canonicalising everything to the homepage de-indexes "
                     "those pages.",
    "host_duplicate": "Pick one canonical host (www or non-www) and 301-redirect the "
                      "other to it so every URL exists only once.",
    "resource_links": "Link to resource files (PDF/CSS/JS/images) only when intended; "
                      "otherwise link to a real HTML page.",
    "empty_anchor": "Give every link descriptive anchor text (or an aria-label) so "
                    "users and search engines understand its destination.",
    "nondesc_anchor": "Replace generic anchor text like \"click here\" / \"read more\" "
                      "with descriptive, keyword-relevant text.",
    "permanent_redirects": "Update internal links to point directly at the final URL "
                           "so requests don't pass through a 301 redirect.",
    "multiple_h1": "Use a single <h1> per page; demote the extra headings to <h2>/<h3>.",
    "multiple_title": "Keep exactly one <title> tag per page; remove the duplicates.",
    "charset": "Declare the character set, e.g. <meta charset=\"utf-8\">, near the top "
               "of <head>.",
    "doctype": "Add <!doctype html> as the first line of every page.",
    "weak_linking": "Add more internal links to weakly-linked pages so they gain "
                    "authority and are easier to discover.",
    "orphans": "Link to orphan pages from relevant pages (nav, related content) so "
               "they are reachable and crawlable.",
    "non_indexable": "Confirm the noindex/robots block is intentional; remove it if the "
                     "page should rank.",
    "canonical": "Add a valid, absolute, self-referencing rel=canonical tag to each "
                 "page.",
    "slow_pages": "Improve load time: optimise images, enable caching/compression, and "
                  "reduce render-blocking resources.",
    "hsts": "Enable HSTS by sending a Strict-Transport-Security response header.",
    "sitemap_robots": "Ensure robots.txt is present with a Sitemap: directive and a "
                      "valid XML sitemap is reachable.",
}


def _doc_issue_rows(start_url, findings, per_issue_cap=60):
    """Flatten findings into report rows (severity, url, issue_type, action),
    ordered errors → warnings → notices. Only checks that actually FAILED are
    included. Each affected URL is its own row, and any "found on" source pages
    are expanded as indented sub-rows so every affected page is listed (capped
    per issue for readability)."""
    order = {"error": 0, "warning": 1, "notice": 2}
    issues = sorted((f for f in findings
                     if f["display"] not in ("passed", "info")),
                    key=lambda f: order.get(f["tier"], 3))
    rows = []
    for f in issues:
        action = RECOMMENDATIONS.get(f["id"], "Review and fix the affected pages.")
        label = TIER_LABEL[f["tier"]]
        items = f["items"]
        if not items:
            rows.append((label, f["tier"], f["summary"], f["name"], action))
            continue
        emitted = 0
        for it in items:
            if emitted >= per_issue_cap:
                break
            note = it.get("note", "")
            url = it.get("url", "")
            cell = f"{url}  —  {note}" if note and note not in url else (url or note)
            rows.append((label, f["tier"], cell, f["name"], action))
            emitted += 1
            for src in it.get("sources", []):
                if emitted >= per_issue_cap:
                    break
                if src == url:
                    continue
                rows.append((label, f["tier"], f"     ↳ found on: {src}",
                             f["name"], ""))
                emitted += 1
        if emitted >= per_issue_cap and len(items) > emitted:
            rows.append((label, f["tier"],
                         f"… and more affected URLs — {f['failed']} total "
                         f"(see Issue Summary for the full count)",
                         f["name"], action))
    return rows


_DOCX_GREEN = (0x16, 0xA3, 0x4A)
_DOCX_AMBER = (0xB4, 0x5F, 0x06)
_DOCX_RED   = (0xC0, 0x27, 0x27)
_DOCX_BLUE  = (0x1D, 0x4E, 0xD8)
_DOCX_GRAY  = (0x64, 0x74, 0x8B)
_DOCX_TIER_RGB = {"error": _DOCX_RED, "warning": _DOCX_AMBER, "notice": _DOCX_BLUE,
                  "info": _DOCX_GRAY, "passed": _DOCX_GREEN}
_DOCX_TIER_FILL = {"error": "FBE9E7", "warning": "FFF3E0", "notice": "E8F0FE",
                   "info": "F1F5F9", "passed": "E8F5E9"}


def _score_rating(score):
    if score >= 90:
        return "Excellent", _DOCX_GREEN
    if score >= 75:
        return "Good", _DOCX_GREEN
    if score >= 60:
        return "Fair", _DOCX_AMBER
    if score >= 40:
        return "Poor", _DOCX_AMBER
    return "Critical", _DOCX_RED


def _docx_shade(cell, fill_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _docx_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Add breathing room inside a table cell (values in twentieths of a point)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _docx_table_borders(table, color="D6DBE3", size=4):
    """Replace the heavy default 'Table Grid' borders with a light, subtle rule."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _docx_no_borders(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _docx_heading(doc, text, color_rgb):
    """A section heading with brand-colour text and a thin rule underneath,
    rather than Word's default plain black Heading style."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_heading("", level=1)
    run = p.add_run(text.upper())
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.size = Pt(15)
    run.font.name = "Calibri Light"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "%02X%02X%02X" % color_rgb)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p


def _docx_page_field(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        run = paragraph.add_run()
        if text is not None:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = text
            run._r.append(it)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), kind)
            run._r.append(fc)


def build_docx(start_url, findings, score, stats, path):
    """Write a complete, client-ready SEO audit report as a Word .docx:
    cover + executive summary + scope/methodology + issue summary + detailed
    findings (Severity / Affected URL / Issue Type / Recommended Action) +
    passed checks + severity definitions, with page numbers."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION

    def C(rgb):
        return RGBColor(*rgb)

    _NAVY = (0x0F, 0x29, 0x42)
    _SLATE = (0x47, 0x55, 0x69)
    _HAIRLINE = "E2E5EA"

    groups = _group(findings)
    n_err, n_warn, n_notice, n_pass = (len(groups["error"]), len(groups["warning"]),
                                       len(groups["notice"]), len(groups["passed"]))
    total_checks = n_err + n_warn + n_notice + n_pass
    rating, rating_rgb = _score_rating(score)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = C((0x22, 0x27, 0x2E))

    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    def styled_table(headers, widths, fill="0F2942", zebra=True):
        t = doc.add_table(rows=1, cols=len(headers))
        t.autofit = True
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _docx_table_borders(t, color=_HAIRLINE)
        for cell, text, w in zip(t.rows[0].cells, headers, widths):
            cell.width = w
            _docx_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _docx_shade(cell, fill)
        if zebra:
            t._zebra_fill = "F4F6F8"
        for row_cells in t.rows:
            for c in row_cells.cells:
                for para in c.paragraphs:
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
        return t

    def add_row_zebra(table, idx):
        cells = table.add_row().cells
        if getattr(table, "_zebra_fill", None) and idx % 2 == 1:
            for c in cells:
                _docx_shade(c, table._zebra_fill)
        for c in cells:
            _docx_cell_margins(c)
            for para in c.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                for r in para.runs:
                    r.font.size = Pt(9.5)
        return cells

    # ---- Footer with page numbers -------------------------------------
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr1 = fp.add_run(f"Technical SEO Audit — {host_key(start_url)}   ·   Page ")
    fr1.font.size = Pt(8)
    fr1.font.color.rgb = C(_SLATE)
    _docx_page_field(fp)

    # ---- Cover: full-width brand banner ---------------------------------
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = True
    bcell = banner.rows[0].cells[0]
    bcell.width = Inches(6.4)
    _docx_shade(bcell, "0F2942")
    _docx_no_borders(banner)
    _docx_cell_margins(bcell, top=520, bottom=520, left=360, right=360)

    p_eyebrow = bcell.paragraphs[0]
    p_eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = p_eyebrow.add_run("TECHNICAL SEO AUDIT REPORT")
    er.bold = True
    er.font.size = Pt(11)
    er.font.color.rgb = C((0x9D, 0xB8, 0xD8))

    p_host = bcell.add_paragraph()
    p_host.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_host.paragraph_format.space_before = Pt(6)
    hr = p_host.add_run(host_key(start_url))
    hr.bold = True
    hr.font.size = Pt(28)
    hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hr.font.name = "Calibri Light"

    p_score = bcell.add_paragraph()
    p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_score.paragraph_format.space_before = Pt(14)
    big = p_score.add_run(f"{score}")
    big.bold = True
    big.font.size = Pt(56)
    big.font.color.rgb = C(rating_rgb)
    big.font.name = "Calibri Light"
    small = p_score.add_run(" / 100")
    small.font.size = Pt(15)
    small.font.color.rgb = C((0xC9, 0xD6, 0xE6))

    p_rating = bcell.add_paragraph()
    p_rating.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rating.paragraph_format.space_before = Pt(2)
    rr = p_rating.add_run(f"SITE HEALTH: {rating.upper()}")
    rr.bold = True
    rr.font.size = Pt(12)
    rr.font.color.rgb = C(rating_rgb)

    p_pass = bcell.add_paragraph()
    p_pass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pass.paragraph_format.space_before = Pt(4)
    pr = p_pass.add_run(
        f"{total_checks - n_err - n_warn - n_notice} of {total_checks} checks passed")
    pr.font.size = Pt(10)
    pr.font.color.rgb = C((0xC9, 0xD6, 0xE6))

    p_counts = bcell.add_paragraph()
    p_counts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_counts.paragraph_format.space_before = Pt(12)
    for i, (tier, n, lbl) in enumerate((("error", n_err, "Errors"),
                                        ("warning", n_warn, "Warnings"),
                                        ("notice", n_notice, "Notices"))):
        if i:
            sep = p_counts.add_run("     ")
        run = p_counts.add_run(f"{n} {lbl}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        dot_color = "%02X%02X%02X" % _DOCX_TIER_RGB[tier]

    doc.add_paragraph()

    # ---- Cover: meta details below the banner --------------------------
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(18)
    meta_run = mp.add_run(
        f"Website: {start_url}\n"
        f"Report date: {stats.get('generated', '')}\n"
        f"Pages crawled: {stats.get('pages', '?')} "
        f"({stats.get('pages_ok', '?')} returned HTTP 200)   ·   "
        f"Links checked: {stats.get('links_checked', '?')}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = C(_SLATE)

    if stats.get("content_warning"):
        wp = doc.add_paragraph()
        wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wp.paragraph_format.space_before = Pt(10)
        wr = wp.add_run(f"⚠ CRITICAL — no indexable content: {stats['content_warning']}")
        wr.bold = True
        wr.font.color.rgb = C(_DOCX_RED)

    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp2.paragraph_format.space_before = Pt(40)
    fr2 = fp2.add_run("Prepared by Canvas Digital")
    fr2.font.size = Pt(9.5)
    fr2.font.color.rgb = C(_SLATE)
    fr2.italic = True

    doc.add_page_break()

    # ---- Executive summary --------------------------------------------
    _docx_heading(doc, "Executive Summary", _NAVY)
    doc.add_paragraph(
        f"This report presents a technical SEO audit of {host_key(start_url)}, "
        f"covering {stats.get('pages', '?')} crawled page(s) evaluated against "
        f"{total_checks} technical checks. The website achieved an overall Site "
        f"Health score of {score}/100 ({rating.lower()}). The audit identified "
        f"{n_err} error-level issue type(s), {n_warn} warning(s) and {n_notice} "
        f"notice(s); {n_pass} check(s) passed with no issues.")

    priorities = [f["name"] for f in groups["error"]] or [f["name"] for f in groups["warning"]]
    if priorities:
        doc.add_paragraph(
            "Priority focus — resolving the following will have the greatest impact:")
        for name in priorities[:6]:
            doc.add_paragraph(name, style="List Bullet")
    else:
        doc.add_paragraph("No critical issues were found. Address the notices below "
                          "for incremental gains.")

    st = styled_table(("Category", "Count", "Meaning"),
                      (Inches(1.4), Inches(0.9), Inches(4.7)))
    for i, (tier, n, meaning) in enumerate((
            ("error", n_err, "Critical — likely to harm indexing or rankings; fix first."),
            ("warning", n_warn, "Moderate — should be addressed to improve SEO health."),
            ("notice", n_notice, "Minor — best-practice recommendations."),
            ("passed", n_pass, "No issues detected for these checks."))):
        cells = add_row_zebra(st, i)
        run = cells[0].paragraphs[0].add_run(TIER_LABEL[tier])
        run.bold = True
        run.font.color.rgb = C(_DOCX_TIER_RGB[tier])
        cells[1].text = str(n)
        cells[2].text = meaning

    # ---- Scope & methodology ------------------------------------------
    _docx_heading(doc, "Scope & Methodology", _NAVY)
    doc.add_paragraph(
        f"Crawl scope: up to {stats.get('pages', '?')} page(s) were crawled starting "
        f"from the homepage and following internal links"
        + ("; the crawl completed fully." if stats.get("crawl_complete")
           else ", stopping at the page cap (some deep pages may be uncrawled).")
        + f" {stats.get('resources_checked', 0)} JavaScript/CSS file(s) were fetched "
        f"and tested for minification. External links "
        + ("were verified." if stats.get("external_checked") else "were not verified."))
    doc.add_paragraph(
        "Site Health scoring: each check contributes a severity-weighted pass ratio "
        "(errors weigh most, then warnings, then notices). The score is the weighted "
        "percentage of checks passed across all crawled pages — the same frequency-"
        "based, severity-weighted methodology used by professional tools such as "
        "Semrush, rather than a flat per-issue penalty.")
    dt = styled_table(("Severity", "Definition"), (Inches(1.4), Inches(5.6)))
    for i, (tier, desc) in enumerate((
            ("error", "A critical problem that can prevent pages from ranking or being "
                      "indexed (e.g. broken links, duplicate content, 4xx pages)."),
            ("warning", "An issue that weakens SEO performance and should be fixed "
                        "(e.g. missing meta descriptions, missing H1, unminified assets)."),
            ("notice", "A minor, best-practice recommendation (e.g. multiple H1s, "
                       "resource-as-page links, HSTS)."))):
        cells = add_row_zebra(dt, i)
        run = cells[0].paragraphs[0].add_run(TIER_LABEL[tier])
        run.bold = True
        run.font.color.rgb = C(_DOCX_TIER_RGB[tier])
        cells[1].text = desc

    # ---- Issue summary (one row per issue type) -----------------------
    doc.add_page_break()
    _docx_heading(doc, "Issue Summary", _NAVY)
    order = {"error": 0, "warning": 1, "notice": 2}
    failing = sorted((f for f in findings
                      if f["display"] not in ("passed", "info")),
                     key=lambda f: order.get(f["tier"], 3))
    if failing:
        summ = styled_table(("Severity", "Issue Type", "Affected", "Recommended Action"),
                            (Inches(0.9), Inches(2.2), Inches(0.8), Inches(3.1)))
        for i, f in enumerate(failing):
            cells = add_row_zebra(summ, i)
            run = cells[0].paragraphs[0].add_run(TIER_LABEL[f["tier"]])
            run.bold = True
            run.font.color.rgb = C(_DOCX_TIER_RGB[f["tier"]])
            cells[1].text = f["name"]
            cells[2].text = f"{f['failed']} {f.get('unit', 'pages')}"
            cells[3].text = RECOMMENDATIONS.get(f["id"], "Review and fix the affected pages.")
    else:
        doc.add_paragraph("No issues found — every check passed. ✓")

    # ---- Detailed findings (Severity / Affected URL / Issue Type / Action)
    doc.add_page_break()
    _docx_heading(doc, "Detailed Findings & Recommended Actions", _NAVY)
    doc.add_paragraph(
        "Each affected URL is listed with its issue type, severity and the "
        "recommended action to resolve it.").italic = True

    rows = _doc_issue_rows(start_url, findings, per_issue_cap=40)
    if rows:
        det = styled_table(("Severity", "Affected URL", "Issue Type", "Recommended Action"),
                           (Inches(0.85), Inches(2.75), Inches(1.7), Inches(2.7)))
        for i, (label, tier, url, issue_type, action) in enumerate(rows):
            cells = add_row_zebra(det, i)
            run = cells[0].paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = C(_DOCX_TIER_RGB.get(tier, (0, 0, 0)))
            _docx_shade(cells[0], _DOCX_TIER_FILL.get(tier, "FFFFFF"))
            cells[1].text = url
            cells[2].text = issue_type
            cells[3].text = action
    else:
        doc.add_paragraph("No issues found — every check passed. ✓")

    # ---- Could not verify (informational, not scored) -----------------
    if groups["info"]:
        doc.add_page_break()
        _docx_heading(doc, "Links That Could Not Be Verified", _SLATE)
        doc.add_paragraph(
            "These links returned a response that neither confirms them working "
            "nor broken — typically because the destination blocks automated "
            "crawlers or was temporarily unavailable. They are NOT counted as "
            "broken and do NOT affect the Site Health score; verify them manually "
            "in a browser.").italic = True
        uv = styled_table(("Link", "Reason", "Found on"),
                          (Inches(3.0), Inches(2.8), Inches(1.2)), fill="475569")
        i = 0
        for f in groups["info"]:
            for it in f["items"]:
                note = it.get("note", "")
                reason = note.split("  (")[0]
                found = ""
                if "linked from" in note:
                    found = note.split("linked from")[-1].strip(" )")
                cells = add_row_zebra(uv, i)
                cells[0].text = it.get("url", "")
                cells[1].text = reason
                cells[2].text = found
                i += 1

    # ---- Passed checks ------------------------------------------------
    if groups["passed"]:
        doc.add_page_break()
        _docx_heading(doc, "Checks Passed", (0x16, 0x7A, 0x3B))
        doc.add_paragraph(
            f"The following {n_pass} check(s) passed with no issues detected.").italic = True
        pt = styled_table(("Check", "Result"), (Inches(3.4), Inches(3.6)), fill="2E7D32")
        for i, f in enumerate(groups["passed"]):
            cells = add_row_zebra(pt, i)
            cells[0].text = f["name"]
            ok = cells[1].paragraphs[0].add_run(f["summary"])
            ok.font.color.rgb = C(_DOCX_GREEN)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Generated by seoaudit.py — Semrush-style Technical SEO Audit. "
        "Site Health = severity-weighted pass ratio across all applicable checks.")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(path)
    return path


def _rtf_escape(s):
    out = []
    for ch in str(s):
        o = ord(ch)
        if ch in "\\{}":
            out.append("\\" + ch)
        elif o > 127:
            out.append(f"\\u{o}?")
        else:
            out.append(ch)
    return "".join(out)


def build_rtf(start_url, findings, score, stats, path):
    """Zero-dependency fallback: write the report as an RTF document (opens in
    Word, Google Docs, LibreOffice) with a colour-coded issue table."""
    groups = _group(findings)
    cf = {"error": 1, "warning": 2, "notice": 3}   # colour table indices
    # Column right-edges in twips (1440/inch): Severity | URL | Issue | Action
    edges = (1200, 5200, 7400, 9360)

    def cell(text, bold=False, color=0):
        pre = ("{" + (r"\cf%d" % color if color else "")
               + (r"\b " if bold else " "))
        return pre + _rtf_escape(text) + r"}\cell "

    def row(cols, bolds, colors):
        r = r"\trowd\trgaph108"
        for e in edges:
            r += r"\cellx%d" % e
        for text, b, c in zip(cols, bolds, colors):
            r += cell(text, b, c)
        return r + r"\row" + "\n"

    parts = [r"{\rtf1\ansi\ansicpg1252\deff0",
             r"{\fonttbl{\f0\fswiss Calibri;}}",
             r"{\colortbl;\red192\green39\blue39;\red180\green95\blue6;"
             r"\red29\green78\blue216;\red22\green163\blue74;}",
             r"\fs32\b Technical SEO Audit Report\b0\par",
             r"\fs22 " + _rtf_escape(start_url) + r"\par",
             r"\fs18 Generated: " + _rtf_escape(stats.get("generated", ""))
             + r"  |  Site Health: " + f"{score}/100"
             + r"  |  " + f"{len(groups['error'])} errors, "
             f"{len(groups['warning'])} warnings, {len(groups['notice'])} notices"
             + r"\par\par",
             r"\fs24\b Issues & Recommended Actions\b0\par\fs18 "]

    # header row
    parts.append(row(("Severity", "Affected URL", "Issue Type", "Recommended Action"),
                     (True, True, True, True), (0, 0, 0, 0)))
    for label, tier, url, issue_type, action in _doc_issue_rows(start_url, findings):
        parts.append(row((label, url, issue_type, action),
                         (True, False, False, False),
                         (cf.get(tier, 0), 0, 0, 0)))
    parts.append(r"\par\i Generated by seoaudit.py — Semrush-style Technical SEO "
                 r"Audit.\i0}")
    with open(path, "w", encoding="ascii", errors="backslashreplace") as fh:
        fh.write("\n".join(parts))
    return path


def write_doc(start_url, findings, score, stats, path):
    """Write the audit report as a document. Uses .docx (python-docx) when the
    path ends in .docx AND the library is available; otherwise writes .rtf.
    Returns the actual path written."""
    want_docx = path.lower().endswith(".docx")
    try:
        import docx  # noqa: F401
        have_docx = True
    except ImportError:
        have_docx = False
    if want_docx and have_docx:
        return build_docx(start_url, findings, score, stats, path)
    # fall back to RTF (swap the extension so the file opens correctly)
    if want_docx and not have_docx:
        path = path[:-5] + ".rtf"
    elif not path.lower().endswith(".rtf"):
        path = path + ".rtf"
    return build_rtf(start_url, findings, score, stats, path)


# ===========================================================================
# Output — JSON
# ===========================================================================
def build_json(start_url, pages, findings, score, stats):
    groups = _group(findings)
    return json.dumps({
        "site": start_url,
        "site_health": score,
        "pages_crawled": len(pages),
        "pages_ok": sum(1 for p in pages.values() if p.ok),
        "links_checked": stats["links_checked"],
        "resources_checked": stats.get("resources_checked", 0),
        "external_checked": stats["external_checked"],
        "rendered": stats.get("rendered", False),
        "content_warning": stats.get("content_warning"),
        "counts": {t: len(groups[t]) for t in ("error", "warning", "notice", "info", "passed")},
        "findings": [{k: v for k, v in f.items()} for f in findings],
    }, indent=2, ensure_ascii=False)


# ===========================================================================
# Main
# ===========================================================================
def main():
    ap = argparse.ArgumentParser(
        description="Automated Technical SEO Audit (Semrush-style).")
    ap.add_argument("url", nargs="?", help="Website to audit (prompted if omitted)")
    ap.add_argument("--max-pages", type=int, default=200,
                    help="Max pages to crawl (default 200). Raise toward the site's "
                         "total page count for full Semrush-style parity.")
    ap.add_argument("--workers", type=int, default=10,
                    help="Concurrent requests (default 10)")
    ap.add_argument("--slow", type=float, default=3.0,
                    help="Slow-page threshold in seconds (default 3.0)")
    ap.add_argument("--max-links", type=int, default=1500,
                    help="Max off-page links to verify (default 1500)")
    ap.add_argument("--max-resources", type=int, default=800,
                    help="Max JS/CSS files to check for minification (default 800)")
    ap.add_argument("--delay", type=float, default=0.0,
                    help="Politeness delay between crawl waves, seconds (default 0)")
    ap.add_argument("--no-external", action="store_true",
                    help="Skip checking external links (faster)")
    ap.add_argument("--no-resources", action="store_true",
                    help="Skip the JS/CSS minification check (faster)")
    ap.add_argument("--render", choices=["auto", "on", "off"], default="auto",
                    help="Render JavaScript with a headless browser (auto/on/off)")
    ap.add_argument("--render-wait", type=int, default=2500,
                    help="ms to wait for JS after load when rendering (default 2500)")
    ap.add_argument("--doc", metavar="FILE",
                    help="Write the audit report as a Word .docx document (this is the "
                         "default deliverable; auto-named if omitted). Falls back to "
                         ".rtf if python-docx is not installed.")
    ap.add_argument("--html", metavar="FILE",
                    help="(Optional) also write an HTML report to FILE")
    ap.add_argument("--json", action="store_true", help="Emit JSON to stdout")
    args = ap.parse_args()

    if args.url:
        url_input = args.url
    else:
        print("\n  Technical SEO Audit — Semrush-style")
        print("  ===================================")
        url_input = input("  Enter website URL (e.g. example.com): ").strip()
        if not url_input:
            sys.exit("  No URL provided. Exiting.")

    start_url = normalize_url(url_input)
    quiet = args.json

    session = requests.Session()
    session.headers.update(BROWSER_HEADERS)

    have_pw = playwright_available()
    render_hint = ""
    if args.render == "off":
        do_render = False
    elif args.render == "on":
        do_render = have_pw
        if not have_pw:
            render_hint = ("--render on requested but Playwright is not installed "
                           "(pip install playwright && playwright install chromium)")
    else:
        seed_probe = fetch_page(session, start_url)
        do_render = have_pw and is_thin(seed_probe)
        if is_thin(seed_probe) and not have_pw:
            render_hint = ("this site is JavaScript-rendered but Playwright is not "
                           "installed, so only the empty HTML shell could be read "
                           "(pip install playwright && playwright install chromium)")

    if do_render:
        if not quiet:
            sys.stderr.write(f"  Rendering {start_url} with headless browser "
                             f"(max {args.max_pages} pages)...\n")
        pages, link_sources, raw_link_sources, crawl_complete = render_crawl(
            start_url, args.max_pages, args.delay, args.render_wait, log=not quiet)
    else:
        if not quiet:
            sys.stderr.write(f"  Crawling {start_url} (max {args.max_pages} pages)...\n")
        pages, link_sources, raw_link_sources, crawl_complete = crawl(
            session, start_url, args.max_pages, args.workers, args.delay,
            log=not quiet)

    if not pages or all(p.error for p in pages.values()):
        first = next(iter(pages.values()), None)
        reason = first.error if first else "no response"
        sys.exit(f"  Could not crawl {start_url}: {reason}")

    robots = fetch_robots(session, start_url)
    sitemap = fetch_sitemaps(session, robots["sitemaps"], start_url)

    # link verification: everything not already crawled as a page
    crawled = set(pages.keys())
    to_check = []
    for target in link_sources:
        if target in crawled:
            continue
        if same_site(start_url, target) or not args.no_external:
            to_check.append(target)
    # Also verify internal links AS WRITTEN when no page was fetched at that
    # exact URL. canon_url() folds www/non-www and trailing-slash variants
    # together, so a link to the non-www homepage resolved to the crawled www
    # page and its 301 was invisible — 154 such links on americanwebbuilders.com.
    crawled_exact = {p.requested_url for p in pages.values()}
    queued = set(to_check)
    for target in raw_link_sources:
        if target in crawled_exact or target in queued:
            continue
        queued.add(target)
        to_check.append(target)
    to_check = to_check[: args.max_links]
    if not quiet and to_check:
        sys.stderr.write(f"  Verifying {len(to_check)} off-page links...\n")
    link_status = check_links(session, to_check, args.workers)

    # resource minification check
    resource_status = {}
    if not args.no_resources:
        assets = set()
        for p in pages.values():
            for a in p.assets:
                if re.search(r"\.(css|js)(\?|#|$)", a, re.I):
                    assets.add(a)
        assets = list(assets)[: args.max_resources]
        if not quiet and assets:
            sys.stderr.write(f"  Checking {len(assets)} JS/CSS files for minification...\n")
        resource_status = check_resources(session, assets, args.workers)

    host_variants = check_host_variants(session, start_url)

    seed_page = pages.get(canon_url(start_url)) or next(iter(pages.values()), None)
    content_warning = None
    if seed_page is not None and is_thin(seed_page):
        content_warning = render_hint or thin_cause(seed_page)

    stats = {
        "links_checked": len(link_status) + len(crawled),
        "resources_checked": len(resource_status),
        "external_checked": not args.no_external,
        "rendered": bool(do_render),
        "content_warning": content_warning,
        "pages": len(pages),
        "pages_ok": sum(1 for p in pages.values() if p.ok),
        "crawl_complete": crawl_complete,
        "generated": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

    findings = analyze(start_url, pages, link_sources, robots, sitemap,
                       link_status, resource_status, args.slow,
                       not args.no_external, crawl_complete, args.max_pages,
                       host_variants=host_variants,
                       raw_link_sources=raw_link_sources)
    score = site_health(findings)
    if content_warning:
        score = min(score, 12)

    if args.json:
        print(build_json(start_url, pages, findings, score, stats))
    else:
        print_console(start_url, pages, findings, score, stats)

    # --- Deliverable: a document (default), NOT HTML ----------------------
    # The audit report is always written as a Word document unless the user
    # asked for JSON only. --doc sets the path; otherwise it is auto-named.
    if not args.json:
        doc_path = args.doc
        if not doc_path:
            host = host_key(start_url) or "site"
            stamp = datetime.datetime.now().strftime("%Y-%m-%d")
            doc_path = f"seo-audit-{host}-{stamp}.docx"
        try:
            written = write_doc(start_url, findings, score, stats, doc_path)
            print(f"\n  Audit report saved to {written}")
        except PermissionError:
            # The target file is open (locked by Word/another app). Write a
            # timestamped copy instead of crashing.
            base, dot, ext = doc_path.rpartition(".")
            alt = f"{base or doc_path}-{datetime.datetime.now():%H%M%S}" \
                  + (f".{ext}" if dot else ".docx")
            written = write_doc(start_url, findings, score, stats, alt)
            print(f"\n  '{doc_path}' is open in another program, so the report was "
                  f"saved to {written} instead.\n  Close the file and re-run to "
                  f"overwrite it.")

    # HTML is only produced when explicitly requested.
    if args.html:
        with open(args.html, "w", encoding="utf-8") as fh:
            fh.write(build_html(start_url, pages, findings, score, stats))
        if not quiet:
            print(f"  HTML report also written to {args.html}")


if __name__ == "__main__":
    main()


